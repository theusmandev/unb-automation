#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
╔══════════════════════════════════════════════════════════════╗
║     Unicode Urdu Novel -> InPage Converter  |  FINAL v3     ║
║                                                              ║
║  Handles:                                                    ║
║  ✓ Urdu text (complete alphabet + diacritics)               ║
║  ✓ English text RTL fix (Congrats Alia! sahi dikhega)       ║
║  ✓ Decorative symbols  ✿ • ★ ❤ etc.                        ║
║  ✓ Line/separator chars ═ ─ ━ ● ◆ etc.                     ║
║  ✓ Smart quotes " " ' '                                     ║
║  ✓ Dashes  – — ...  (en dash, em dash, ellipsis)            ║
║  ✓ Urdu digits  ۰۱۲۳۴۵۶۷۸۹                                  ║
║  ✓ Mixed Urdu+English lines                                  ║
║  ✓ .txt aur .docx dono formats                              ║
║  ✓ Full report: kya convert hua, kya nahi                   ║
╚══════════════════════════════════════════════════════════════╝

  INSTALL (sirf docx ke liye zaroori):
      pip install python-docx

  USE:
      1. Neeche INPUT_PATH aur OUTPUT_PATH set karo
      2. SYMBOL_ACTION choose karo
      3. python urdu_novel_to_inpage_FINAL.py
"""

import sys
import os
import re
from collections import Counter

# ══════════════════════════════════════════════════════════════
#  ⚙️  SETTINGS — Yahan sirf yahi section change karna hai
# ══════════════════════════════════════════════════════════════

INPUT_PATH  = r"C:\Users\PCS\Downloads\Qafas\قفس قسط نمبر 8.docx"
OUTPUT_PATH = r"C:\Users\PCS\Downloads\Qafas\قفس قسط نمبر 8-done.txt"

# Decorative symbols (✿ • ★ etc.) ka kya karna hai:
#   'replace'  ->  neeche SYMBOL_REPLACEMENTS mein jo likha hai woh lagao
#   'remove'   ->  bilkul hata do
#   'keep'     ->  jaise hai waise rakho (InPage mein ? ya box dikhe ga)
SYMBOL_ACTION = 'replace'

# Agar SYMBOL_ACTION = 'replace' hai to har symbol ki jagah kya aaye
# Aap yahan apni marzi se koi bhi character likh sakte hain
SYMBOL_REPLACEMENTS = {
    # Decorative / floral
    '✿': '🕸️', '❀': '*', '✾': '*', '❁': '*', '✽': '*',
    '❃': '*', '✸': '*', '✺': '*', '✻': '*', '✼': '*',
    '❋': '*',
    # Bullets / dots
    '•': '-', '·': '.', '‧': '.', '∙': '.',
    '●': 'o', '◉': 'o', '○': 'o', '◎': 'o',
    '◆': '*', '◇': '*', '■': '*', '□': '*',
    '▪': '-', '▫': '-', '▸': '>', '▹': '>',
    # Stars / hearts
    '★': '*', '☆': '*', '✦': '*', '✧': '*', '✩': '*',
    '✪': '*', '✫': '*', '✬': '*', '✭': '*', '✮': '*',
    '❤': '<3', '♥': '<3', '♡': '<3', '❥': '<3',
    '☀': '*', '✨': '*', '💫': '*', '🌟': '*',
    # Box-drawing / separators
    '═': '=', '─': '-', '━': '-', '│': '|', '┃': '|',
    '╔': '+', '╗': '+', '╚': '+', '╝': '+', '╠': '+',
    '╣': '+', '╦': '+', '╩': '+', '╬': '+',
    '▬': '-', '▭': '-', '▰': '-',
    # Typography
    '…': '...',   # ellipsis -> 3 dots
    '–': '-',     # en dash
    '—': '-',     # em dash
    '―': '-',     # horizontal bar
    '\u00AB': '"', '\u00BB': '"',  # « »
    '\u2039': '',  # ‹  (already used in Urdu mapping, skip)
    '\u203A': '',  # ›  (already used in Urdu mapping, skip)
    # Misc
    '©': '(c)', '®': '(r)', '™': '(tm)',
    '°': 'deg', '±': '+/-',
    '½': '1/2', '¼': '1/4', '¾': '3/4',
    '\u200B': '',  # zero-width space -> remove
    '\u200C': '',  # zero-width non-joiner -> remove
    '\u200D': '',  # zero-width joiner -> remove
    '\uFEFF': '',  # BOM -> remove
    '\u00A0': ' ', # non-breaking space -> normal space
}

# ══════════════════════════════════════════════════════════════
#  INTERNAL CONSTANTS  (change mat karna)
# ══════════════════════════════════════════════════════════════

CTRL = '\u0004'

PRE_REPLACEMENTS = [
    ('\u06CC\u0626', CTRL + '\u00A4' + CTRL + '\u00BF'),
    ('\u0623',        CTRL + '\u0081' + CTRL + '\u00BF'),
    ('\u0624',        CTRL + '\u00A2' + CTRL + '\u00BF'),
    ('\u0622',        CTRL + '\u0081' + CTRL + '\u00B3'),
]

_PRE_REPLACEMENT_CHARS = {'\u06CC', '\u0626', '\u0623', '\u0624', '\u0622'}

UNICODE_TO_INPAGE = {
    '\u06BA': CTRL + '\u00A1', '\u064D': CTRL + '\u00A8',
    '\u06F1': CTRL + '\u00D1', '\u06F2': CTRL + '\u00D2', '\u06F3': CTRL + '\u00D3',
    '\u06F4': CTRL + '\u00D4', '\u06F5': CTRL + '\u00D5', '\u06F6': CTRL + '\u00D6',
    '\u06F7': CTRL + '\u00D7', '\u06F8': CTRL + '\u00D8', '\u06F9': CTRL + '\u00D9',
    '\u06F0': CTRL + '\u00D0',
    '\u0642': CTRL + '\u203A', '\u0648': CTRL + '\u00A2', '\u0639': CTRL + '\u02DC',
    '\u0631': CTRL + '\u017D', '\u062A': CTRL + '\u201E', '\u06D2': CTRL + '\u00A5',
    '\u0621': CTRL + '\u00A3', '\u06CC': CTRL + '\u00A4', '\u06C1': CTRL + '\u00A6',
    '\u067E': CTRL + '\u0192', '\u060E': CTRL + '\u00F2', '\u0627': CTRL + '\u0081',
    '\u0633': CTRL + '\u2019', '\u062F': CTRL + '\u2039', '\u0641': CTRL + '\u0161',
    '\u06AF': CTRL + '\u009D', '\u06BE': CTRL + '\u00A7', '\u062C': CTRL + '\u2021',
    '\u06A9': CTRL + '\u0153', '\u0644': CTRL + '\u017E', '\u061B': CTRL + '\u00EA',
    '\u0632': CTRL + '\u0090', '\u0634': CTRL + '\u201C', '\u0686': CTRL + '\u02C6',
    '\u0637': CTRL + '\u2013', '\u0628': CTRL + '\u201A', '\u0646': CTRL + '\u00A0',
    '\u0645': CTRL + '\u0178', '\u060C': CTRL + '\u00ED', '\u06D4': CTRL + '\u00F3',
    '\u064B': CTRL + '\u00C7', '\u064A': CTRL + '\u00B8', '\u0610': CTRL + '\u00F8',
    '\u0626': CTRL + '\u00A3', '\u064C': CTRL + '\u00B5', '\u0651': CTRL + '\u00AD',
    '\u0652': CTRL + '\u00B1', '\uFDFA': CTRL + '\u00F6', '\u0611': CTRL + '\u00AE',
    '\u0691': CTRL + '\u008F', '\u0679': CTRL + '\u2026', '\u0601': CTRL + '\u00F7',
    '\u0657': CTRL + '\u00BE', '\u0670': CTRL + '\u00BD', '\u06C3': CTRL + '\u00B9',
    '\u064F': CTRL + '\u00AC', '\u2018': CTRL + '\u00FD', '\u2019': CTRL + '\u00FE',
    '\u0614': CTRL + '\u00CF', '\u0653': CTRL + '\u00B3', '\u0635': CTRL + '\u201D',
    '\u0688': CTRL + '\u0152', '\u0656': CTRL + '\u00B0', '\u063A': CTRL + '\u2122',
    '\u062D': CTRL + '\u2030', '\u0636': CTRL + '\u2022', '\u062E': CTRL + '\u0160',
    '\u0612': CTRL + '\u00E7', '\u0630': CTRL + '\u008D', '\u0698': CTRL + '\u2018',
    '\u062B': CTRL + '\u2020', '\u0638': CTRL + '\u2014', '\u0613': CTRL + '\u00E6',
    '\u0650': CTRL + '\u00AA', '\u064E': CTRL + '\u00AB', '\u061F': CTRL + '\u00EE',
    '[': CTRL + '\u00FA', ']': CTRL + '\u00FB', '.': CTRL + '\u00FC',
    '!': CTRL + '\u00DA', ',': CTRL + '\u00F9', '/': CTRL + '\u00F1',
    ')': CTRL + '\u00E1', '(': CTRL + '\u00E2', ':': CTRL + '\u00E9',
    ' ': CTRL + ' ',
}

_RECOGNIZED_CHARS = set(UNICODE_TO_INPAGE.keys()) | _PRE_REPLACEMENT_CHARS
_URDU_RE   = re.compile(r'[\u0600-\u06FF\u0750-\u077F\uFB50-\uFDFF\uFE70-\uFEFF]')
_LATIN_RE  = re.compile(r'[A-Za-z]')


# ══════════════════════════════════════════════════════════════
#  STEP 1: Symbol Cleanup
# ══════════════════════════════════════════════════════════════

def apply_symbol_replacements(text: str) -> tuple[str, Counter]:
    """
    SYMBOL_REPLACEMENTS apply karo.
    Returns: (cleaned_text, counter_of_what_was_replaced)
    """
    replaced_counts = Counter()
    for sym, replacement in SYMBOL_REPLACEMENTS.items():
        if sym in text:
            count = text.count(sym)
            replaced_counts[sym] += count
            text = text.replace(sym, replacement)
    return text, replaced_counts


# ══════════════════════════════════════════════════════════════
#  STEP 2: English RTL Fix
# ══════════════════════════════════════════════════════════════

def fix_english_for_rtl(line: str) -> str:
    """
    InPage RTL mode mein English (LTR) text ulta ho jata hai.
    Is function mein English segments detect ho ke unhe pre-reverse
    kiya jata hai taake InPage mein sahi dikhe.

    Misal:
        "Congrats Alia!"  ->  stores as "!Alia Congrats"
        InPage render:    ->  shows    "Congrats Alia!"  ✅
    """
    if not _LATIN_RE.search(line):
        return line  # sirf Urdu, kuch nahi karna

    # Line ko segments mein todo: urdu ya english
    segments = []
    current = ''
    current_type = None

    for ch in line:
        if _URDU_RE.match(ch):
            ch_type = 'urdu'
        elif _LATIN_RE.match(ch) or ch.isdigit():
            ch_type = 'english'
        else:
            ch_type = current_type or 'other'

        if ch_type != current_type and current:
            segments.append((current_type, current))
            current = ch
            current_type = ch_type
        else:
            current += ch
            current_type = ch_type

    if current:
        segments.append((current_type, current))

    result = []
    for seg_type, seg_text in segments:
        if seg_type == 'english' and _LATIN_RE.search(seg_text):
            stripped   = seg_text.strip()
            lead_sp    = seg_text[:len(seg_text) - len(seg_text.lstrip())]
            trail_sp   = seg_text[len(seg_text.rstrip()):]

            # leading/trailing punctuation alag karo
            lead_p = ''
            trail_p = ''
            tmp = stripped
            while tmp and not tmp[0].isalnum() and tmp[0] != ' ':
                lead_p += tmp[0]; tmp = tmp[1:]
            while tmp and not tmp[-1].isalnum() and tmp[-1] != ' ':
                trail_p = tmp[-1] + trail_p; tmp = tmp[:-1]

            words = tmp.split(' ')
            words.reverse()
            reversed_text = trail_p + ' '.join(words) + lead_p
            result.append(lead_sp + reversed_text + trail_sp)
        else:
            result.append(seg_text)

    return ''.join(result)


# ══════════════════════════════════════════════════════════════
#  STEP 3: Core Unicode -> InPage Conversion
# ══════════════════════════════════════════════════════════════

def convert_line(line: str) -> tuple[str, set]:
    """Single line convert karo. Returns (converted, unmapped_chars)"""
    unmapped = set()

    for ch in line:
        if ch in ('\n', '\t', '\r', CTRL):
            continue
        if ch.strip() and ch not in _RECOGNIZED_CHARS:
            unmapped.add(ch)

    # Pre-replacements (multi-char Urdu combos)
    for old, new in PRE_REPLACEMENTS:
        line = line.replace(old, new)

    # Char-by-char mapping
    result = []
    for ch in line:
        if ch in UNICODE_TO_INPAGE:
            result.append(UNICODE_TO_INPAGE[ch])
        elif ch in ('\n', '\r', '\t'):
            result.append(ch)
        elif ch == CTRL:
            result.append(ch)  # already-converted output preserve karo
        else:
            result.append(ch)  # unmapped: as-is

    return ''.join(result), unmapped


def convert_text(text: str) -> tuple[str, set, int, int]:
    """
    Pura text convert karo.
    Returns: (converted_text, all_unmapped, total_chars, converted_chars)
    """
    lines = text.split('\n')
    converted_lines = []
    all_unmapped = set()
    total_chars = 0
    unconverted_chars = 0

    for line in lines:
        # RTL fix
        line = fix_english_for_rtl(line)
        # Convert
        converted, unmapped = convert_line(line)
        converted_lines.append(converted)
        all_unmapped |= unmapped

        for ch in line:
            if ch not in ('\n', '\r', '\t') and ch.strip():
                total_chars += 1
                if ch in unmapped:
                    unconverted_chars += 1

    return '\n'.join(converted_lines), all_unmapped, total_chars, unconverted_chars


# ══════════════════════════════════════════════════════════════
#  FILE READERS
# ══════════════════════════════════════════════════════════════

def read_txt(path: str) -> str:
    for enc in ['utf-8-sig', 'utf-8', 'utf-16', 'utf-16-le', 'utf-16-be', 'cp1256', 'iso-8859-6']:
        try:
            with open(path, 'r', encoding=enc) as f:
                content = f.read()
            print(f"    ✓ Encoding: {enc}")
            return content
        except (UnicodeDecodeError, LookupError):
            continue
    raise ValueError(f"File '{path}' kisi encoding se nahi parhi ja saki.")


def read_docx(path: str) -> str:
    try:
        from docx import Document
    except ImportError:
        print("\n  ✗ python-docx install karo:  pip install python-docx\n")
        sys.exit(1)
    doc = Document(path)
    lines = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            cells = ' | '.join(c.text.strip() for c in row.cells if c.text.strip())
            if cells:
                lines.append(cells)
    return '\n'.join(lines)


# ══════════════════════════════════════════════════════════════
#  REPORT PRINTER
# ══════════════════════════════════════════════════════════════

def print_report(unmapped: set, replaced: Counter, total: int, unconverted: int):
    converted = total - unconverted
    pct = round(converted / max(total, 1) * 100, 2)

    print()
    print("┌─────────────────────────────────────────────┐")
    print("│              CONVERSION REPORT               │")
    print("├─────────────────────────────────────────────┤")
    print(f"│  Total characters    : {total:>8,}             │")
    print(f"│  Successfully mapped : {converted:>8,}  ({pct}%)  │")
    print(f"│  Unmapped (as-is)    : {unconverted:>8,}             │")
    print("├─────────────────────────────────────────────┤")

    if replaced:
        print("│  Symbols replaced:                          │")
        for sym, cnt in replaced.most_common(10):
            name = f"U+{ord(sym):04X}"
            print(f"│    '{sym}' ({name}) x{cnt:<6}                 │")

    if unmapped:
        print("├─────────────────────────────────────────────┤")
        print("│  Still unmapped (check InPage output):      │")
        for ch in sorted(unmapped)[:15]:
            print(f"│    '{ch}'  U+{ord(ch):04X}                        │")
        if len(unmapped) > 15:
            print(f"│    ... aur {len(unmapped)-15} aur characters          │")
    else:
        print("│  ✓ Koi unmapped character nahi!             │")

    print("└─────────────────────────────────────────────┘")


# ══════════════════════════════════════════════════════════════
#  MAIN
# ══════════════════════════════════════════════════════════════

def main():
    print()
    print("╔══════════════════════════════════════════════════════╗")
    print("║    Urdu Novel -> InPage Converter  |  FINAL v3      ║")
    print("╚══════════════════════════════════════════════════════╝")
    print(f"  Input   : {INPUT_PATH}")
    print(f"  Output  : {OUTPUT_PATH}")
    print(f"  Symbols : {SYMBOL_ACTION}")
    print("─" * 56)

    # ── Checks ──────────────────────────────────────────────
    if not os.path.exists(INPUT_PATH):
        print(f"\n  ✗ File nahi mili: {INPUT_PATH}")
        sys.exit(1)

    ext = os.path.splitext(INPUT_PATH)[1].lower()
    if ext not in ('.txt', '.docx'):
        print("\n  ✗ Sirf .txt aur .docx support hain.")
        sys.exit(1)

    # ── Read ────────────────────────────────────────────────
    print("\n  [1/4] File parh raha hoon...")
    text = read_txt(INPUT_PATH) if ext == '.txt' else read_docx(INPUT_PATH)
    char_count = len(text)
    word_count = len(text.split())
    line_count = text.count('\n') + 1
    print(f"    {line_count:,} lines | {word_count:,} words | {char_count:,} chars")

    # ── Symbol cleanup ──────────────────────────────────────
    print("\n  [2/4] Symbols clean kar raha hoon...")
    if SYMBOL_ACTION == 'replace':
        text, replaced_counts = apply_symbol_replacements(text)
        print(f"    {sum(replaced_counts.values()):,} symbols replace hue ({len(replaced_counts)} unique types)")
    elif SYMBOL_ACTION == 'remove':
        replaced_counts = Counter()
        for sym in SYMBOL_REPLACEMENTS:
            if sym in text:
                replaced_counts[sym] += text.count(sym)
                text = text.replace(sym, '')
        print(f"    {sum(replaced_counts.values()):,} symbols remove hue")
    else:  # keep
        replaced_counts = Counter()
        print("    Symbols as-is rakhe gaye")

    # ── Convert ─────────────────────────────────────────────
    print("\n  [3/4] Convert + English RTL fix ho raha hai...")
    converted, unmapped, total_chars, unconverted_chars = convert_text(text)
    pct = round((total_chars - unconverted_chars) / max(total_chars, 1) * 100, 1)
    print(f"    {pct}% characters successfully convert hue")

    # ── Save ────────────────────────────────────────────────
    print("\n  [4/4] File save ho rahi hai...")
    os.makedirs(os.path.dirname(OUTPUT_PATH) or '.', exist_ok=True)
    with open(OUTPUT_PATH, 'w', encoding='utf-8-sig', newline='') as f:
        f.write(converted)
    size_kb = os.path.getsize(OUTPUT_PATH) / 1024
    print(f"    ✓ '{os.path.basename(OUTPUT_PATH)}'  ({size_kb:.1f} KB)")

    # ── Report ──────────────────────────────────────────────
    print_report(unmapped, replaced_counts, total_chars, unconverted_chars)

    print()
    print("╔══════════════════════════════════════════════════════╗")
    print("║              INPAGE MEIN USE KARNE KA TARIKA        ║")
    print("╠══════════════════════════════════════════════════════╣")
    print(f"║  1. '{os.path.basename(OUTPUT_PATH)}'               ")
    print("║     Notepad mein kholo                               ║")
    print("║  2. Ctrl+A -> Ctrl+C  (sab copy karo)               ║")
    print("║  3. InPage mein Text Box banao                       ║")
    print("║  4. Ctrl+V paste karo                               ║")
    print("║  5. Font: Noori Nastaleeq ya Jameel Noori            ║")
    print("║  6. Direction: RTL confirm karo                      ║")
    print("║  ─────────────────────────────────────────────────  ║")
    print("║  TIP: Pehle sirf ek chapter paste karo aur check    ║")
    print("║       karo. Theek hai to puri novel karo.           ║")
    print("╚══════════════════════════════════════════════════════╝")
    print()


if __name__ == '__main__':
    main()
