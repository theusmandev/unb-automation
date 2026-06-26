#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
==============================================================
  Unicode Urdu -> InPage Converter  |  v2 (Fixed)
==============================================================
  Fixes in this version:
  1. English text (LTR) ka order reverse kiya gaya hai taake
     InPage RTL mode mein sahi dikhe.
     Misal: "Congrats Alia!" -> InPage mein "!Alia Congrats"
     tha, ab sahi "Congrats Alia!" dikhega.

  2. Decorative Unicode symbols jaise ✿ ★ • etc. jo InPage
     mein ? ban jaate the -- ab user choose kar sakta hai:
     inhe rakhna ya remove karna.
==============================================================
"""

import sys
import os
import re

CTRL = '\u0004'

# ─────────────────────────────────────────────────────────────────────────────
#  SETTING: Decorative/unmapped symbols ka kya karna hai?
#    'keep'   -> as-is rakho (InPage mein ? ya box dikhega)
#    'remove' -> bilkul hata do
# ─────────────────────────────────────────────────────────────────────────────
UNMAPPED_SYMBOL_ACTION = 'keep'  # <-- yahan 'remove' likh do agar hatana ho

# ─────────────────────────────────────────────────────────────────────────────
#  PRE-REPLACEMENTS (multi-char sequences)
# ─────────────────────────────────────────────────────────────────────────────
PRE_REPLACEMENTS = [
    ('\u06CC\u0626', CTRL + '\u00A4' + CTRL + '\u00BF'),
    ('\u0623',        CTRL + '\u0081' + CTRL + '\u00BF'),
    ('\u0624',        CTRL + '\u00A2' + CTRL + '\u00BF'),
    ('\u0622',        CTRL + '\u0081' + CTRL + '\u00B3'),
]

_PRE_REPLACEMENT_CHARS = {'\u06CC', '\u0626', '\u0623', '\u0624', '\u0622'}

# ─────────────────────────────────────────────────────────────────────────────
#  Mapping Table
# ─────────────────────────────────────────────────────────────────────────────
UNICODE_TO_INPAGE = {
    '\u06BA': CTRL + '\u00A1', '\u064D': CTRL + '\u00A8',
    '\u06F1': CTRL + '\u00D1', '\u06F2': CTRL + '\u00D2', '\u06F3': CTRL + '\u00D3',
    '\u06F4': CTRL + '\u00D4', '\u06F5': CTRL + '\u00D5', '\u06F6': CTRL + '\u00D6',
    '\u06F7': CTRL + '\u00D7', '\u06F8': CTRL + '\u00D8', '\u06F9': CTRL + '\u00D9',
    '\u06F0': CTRL + '\u00D0',
    '\u0642': CTRL + '\u203A', '\u0648': CTRL + '\u00A2', '\u0639': CTRL + '\u02DC',
    '\u0631': CTRL + '\u017D', '\u062A': CTRL + '\u201E', '\u06D2': CTRL + '\u00A5',
    '\u0621': CTRL + '\u00A3', '\u06CC': CTRL + '\u00A4', '\u06C1': CTRL + '\u00A6',
    '\u067E': CTRL + '\u0192', '\u060E': CTRL + '\u00F2', '\u0627': CTRL + '\u0081',
    '\u0633': CTRL + '\u2019', '\u062F': CTRL + '\u2039', '\u0641': CTRL + '\u0161',
    '\u06AF': CTRL + '\u009D', '\u06BE': CTRL + '\u00A7', '\u062C': CTRL + '\u2021',
    '\u06A9': CTRL + '\u0153', '\u0644': CTRL + '\u017E', '\u061B': CTRL + '\u00EA',
    '\u0632': CTRL + '\u0090', '\u0634': CTRL + '\u201C', '\u0686': CTRL + '\u02C6',
    '\u0637': CTRL + '\u2013', '\u0628': CTRL + '\u201A', '\u0646': CTRL + '\u00A0',
    '\u0645': CTRL + '\u0178', '\u060C': CTRL + '\u00ED', '\u06D4': CTRL + '\u00F3',
    '\u064B': CTRL + '\u00C7', '\u064A': CTRL + '\u00B8', '\u0610': CTRL + '\u00F8',
    '\u0626': CTRL + '\u00A3', '\u064C': CTRL + '\u00B5', '\u0651': CTRL + '\u00AD',
    '\u0652': CTRL + '\u00B1', '\uFDFA': CTRL + '\u00F6', '\u0611': CTRL + '\u00AE',
    '\u0691': CTRL + '\u008F', '\u0679': CTRL + '\u2026', '\u0601': CTRL + '\u00F7',
    '\u0657': CTRL + '\u00BE', '\u0670': CTRL + '\u00BD', '\u06C3': CTRL + '\u00B9',
    '\u064F': CTRL + '\u00AC', '\u2018': CTRL + '\u00FD', '\u2019': CTRL + '\u00FE',
    '\u0614': CTRL + '\u00CF', '\u0653': CTRL + '\u00B3', '\u0635': CTRL + '\u201D',
    '\u0688': CTRL + '\u0152', '\u0656': CTRL + '\u00B0', '\u063A': CTRL + '\u2122',
    '\u062D': CTRL + '\u2030', '\u0636': CTRL + '\u2022', '\u062E': CTRL + '\u0160',
    '\u0612': CTRL + '\u00E7', '\u0630': CTRL + '\u008D', '\u0698': CTRL + '\u2018',
    '\u062B': CTRL + '\u2020', '\u0638': CTRL + '\u2014', '\u0613': CTRL + '\u00E6',
    '\u0650': CTRL + '\u00AA', '\u064E': CTRL + '\u00AB', '\u061F': CTRL + '\u00EE',
    '[': CTRL + '\u00FA', ']': CTRL + '\u00FB', '.': CTRL + '\u00FC',
    '!': CTRL + '\u00DA', ',': CTRL + '\u00F9', '/': CTRL + '\u00F1',
    ')': CTRL + '\u00E1', '(': CTRL + '\u00E2', ':': CTRL + '\u00E9',
    ' ': CTRL + ' ',
}

_RECOGNIZED_CHARS = set(UNICODE_TO_INPAGE.keys()) | _PRE_REPLACEMENT_CHARS


# ─────────────────────────────────────────────────────────────────────────────
#  English Text Reversal Fix
#
#  InPage RTL mode mein jab English (LTR) text hota hai to uska character
#  order ulta ho jata hai. Is fix mein hum English words/sentences ko
#  detect karte hain aur unhe reverse karte hain taake InPage mein sahi dikhe.
#
#  Approach:
#  - Ek line ko "segments" mein todo: Urdu segments aur English segments
#  - English segment = consecutive English letters, digits, spaces, punctuation
#    jo Urdu nahi hain
#  - English segment ke tokens (words + punctuation) ko reverse karo
# ─────────────────────────────────────────────────────────────────────────────

# Urdu/Arabic Unicode range
_URDU_RE = re.compile(r'[\u0600-\u06FF\u0750-\u077F\uFB50-\uFDFF\uFE70-\uFEFF]')

# English segment: ASCII letters, digits, basic punctuation, spaces
# (CTRL char \u0004 bhi preserve karna hai)
_ENG_SEGMENT_RE = re.compile(r'([A-Za-z0-9 \t\'"\-_@#$%^&*+=<>;|~`]+(?:[!?,.:]+[A-Za-z0-9 \t\'"\-_@#$%^&*+=<>;|~`]*)*)')


def _has_urdu(s):
    return bool(_URDU_RE.search(s))


def _has_english(s):
    return bool(re.search(r'[A-Za-z]', s))


def fix_english_for_rtl(line: str) -> str:
    """
    Ek line mein English segments dhundho aur unhe reverse karo
    taake InPage RTL mein sahi order mein dikhe.

    Misal:
      Input:  "Congrats Alia!"
      Output: "!Alia Congrats"   <-- InPage is ulte version ko RTL
                                      mein sahi "Congrats Alia!" dikhayega
    """
    # Agar line mein sirf English hai (koi Urdu nahi) to reverse karo
    # Agar mixed hai to English tokens ko reverse karo

    # Line ko split karo: Urdu parts aur English parts
    # Strategy: word-by-word tokenize, English consecutive words ko group karo
    # aur reverse karo

    # Simple approach: pure-English lines ko word-level reverse karo
    # Mixed lines: English word groups ko reverse karo

    if not _has_english(line):
        return line  # sirf Urdu, kuch nahi karna

    # Tokens mein todo: space-separated nahi, character by character grouping
    # Hum segments banate hain: [urdu_segment, english_segment, ...]
    segments = []
    i = 0
    current = ''
    current_type = None  # 'urdu' ya 'english' ya 'other'

    for ch in line:
        if _URDU_RE.match(ch):
            ch_type = 'urdu'
        elif re.match(r'[A-Za-z0-9]', ch):
            ch_type = 'english'
        else:
            # space, punctuation: previous type ke sath rakho
            ch_type = current_type if current_type else 'other'

        if ch_type != current_type and current:
            segments.append((current_type, current))
            current = ch
            current_type = ch_type
        else:
            current += ch
            current_type = ch_type

    if current:
        segments.append((current_type, current))

    # Ab English segments ko process karo
    result = []
    for seg_type, seg_text in segments:
        if seg_type == 'english' and _has_english(seg_text):
            # Words aur punctuation ko reverse karo
            # "Congrats Alia!" -> tokens: ["Congrats", " ", "Alia", "!"]
            # Reverse: ["!", "Alia", " ", "Congrats"]
            # Lekin punctuation apne word ke sath rakho

            # Approach: strip trailing punctuation, reverse words, re-attach punctuation
            stripped = seg_text.strip()
            leading_space = seg_text[:len(seg_text) - len(seg_text.lstrip())]
            trailing_space = seg_text[len(seg_text.rstrip()):]

            # Punctuation at start/end
            leading_punct = ''
            trailing_punct = ''
            temp = stripped
            while temp and not temp[0].isalnum() and temp[0] != ' ':
                leading_punct += temp[0]
                temp = temp[1:]
            while temp and not temp[-1].isalnum() and temp[-1] != ' ':
                trailing_punct = temp[-1] + trailing_punct
                temp = temp[:-1]

            # Words reverse karo
            words = temp.split(' ')
            words.reverse()
            reversed_text = trailing_punct + ' '.join(words) + leading_punct
            result.append(leading_space + reversed_text + trailing_space)
        else:
            result.append(seg_text)

    return ''.join(result)


# ─────────────────────────────────────────────────────────────────────────────
#  Core Conversion
# ─────────────────────────────────────────────────────────────────────────────
def convert_unicode_to_inpage(text: str) -> tuple:
    unmapped = set()
    for ch in text:
        if ch in ('\n', '\t', '\r'):
            continue
        if ch.strip() and ch not in _RECOGNIZED_CHARS:
            unmapped.add(ch)

    # English RTL fix: line by line
    lines = text.split('\n')
    fixed_lines = []
    for line in lines:
        fixed_lines.append(fix_english_for_rtl(line))
    text = '\n'.join(fixed_lines)

    # Pre-replacements
    for old, new in PRE_REPLACEMENTS:
        text = text.replace(old, new)

    # Character mapping
    result = []
    for ch in text:
        if ch in UNICODE_TO_INPAGE:
            result.append(UNICODE_TO_INPAGE[ch])
        elif ch in ('\n', '\t', '\r'):
            result.append(ch)
        elif ch in unmapped and UNMAPPED_SYMBOL_ACTION == 'remove':
            pass  # hata do
        else:
            result.append(ch)  # as-is rakho

    return ''.join(result), unmapped


# ─────────────────────────────────────────────────────────────────────────────
#  File Readers
# ─────────────────────────────────────────────────────────────────────────────
def read_txt_file(path: str) -> str:
    for enc in ['utf-8-sig', 'utf-8', 'utf-16', 'cp1256', 'iso-8859-6']:
        try:
            with open(path, 'r', encoding=enc) as f:
                content = f.read()
            print(f"    Encoding detect hua: {enc}")
            return content
        except (UnicodeDecodeError, LookupError):
            continue
    raise ValueError(f"File '{path}' ko kisi bhi encoding se nahi parh saka.")


def read_docx_file(path: str) -> str:
    try:
        from docx import Document
    except ImportError:
        print("\n  X python-docx install karo:")
        print("     pip install python-docx\n")
        sys.exit(1)

    doc = Document(path)
    lines = []
    for para in doc.paragraphs:
        lines.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                lines.append(row_text)
    return '\n'.join(lines)


# ─────────────────────────────────────────────────────────────────────────────
#  Main
# ─────────────────────────────────────────────────────────────────────────────
def main():
    # 👇 Apne paths yahan set karein 👇
    input_path  = r"C:\Users\PCS\Downloads\temp\New Text Document.txt"
    output_path = r"C:\Users\PCS\Downloads\temp\New Text Document_done.txt"

    if not os.path.exists(input_path):
        print(f"\nX File nahi mili: '{input_path}'")
        sys.exit(1)

    ext = os.path.splitext(input_path)[1].lower()
    if ext not in ('.txt', '.docx'):
        print("\nX Sirf .txt aur .docx files support hain.")
        sys.exit(1)

    print()
    print("=" * 60)
    print("   Unicode Urdu -> InPage Converter  (v2 - Fixed)")
    print("=" * 60)
    print(f"  Input   : {input_path}")
    print(f"  Output  : {output_path}")
    print(f"  Symbols : {UNMAPPED_SYMBOL_ACTION}  (✿ jaise symbols)")
    print("-" * 60)

    print("\n  File parh raha hoon...")
    try:
        text = read_txt_file(input_path) if ext == '.txt' else read_docx_file(input_path)
    except Exception as e:
        print(f"\n  X Error: {e}")
        sys.exit(1)

    word_count = len(text.split())
    char_count = len(text)
    print(f"    {char_count:,} characters | ~{word_count:,} words mile")

    print("\n  English RTL fix + Convert ho raha hai...")
    converted, unmapped = convert_unicode_to_inpage(text)

    pct = round(100 - (len(unmapped) / max(char_count, 1) * 100), 1)
    print(f"    {pct}% characters successfully convert hue")

    if unmapped:
        sample = ', '.join(f"'{c}'(U+{ord(c):04X})" for c in sorted(unmapped)[:10])
        action_msg = "hata diye gaye" if UNMAPPED_SYMBOL_ACTION == 'remove' else "as-is rakhe gaye"
        print(f"    Unmapped symbols ({action_msg}): {sample}")
        print(f"    (Agar inhe hatana ho to script ke upar UNMAPPED_SYMBOL_ACTION = 'remove' set karo)")

    print("\n  File save ho rahi hai...")
    try:
        with open(output_path, 'w', encoding='utf-8-sig', newline='') as f:
            f.write(converted)
    except Exception as e:
        print(f"\n  X Save error: {e}")
        sys.exit(1)

    size_kb = os.path.getsize(output_path) / 1024
    print(f"    '{os.path.basename(output_path)}' save ho gayi! ({size_kb:.1f} KB)")

    print()
    print("=" * 60)
    print("  KAAM MUKAMMAL HO GAYA!")
    print("-" * 60)
    print("  InPage mein use karne ka tarika:")
    print(f"  1. '{os.path.basename(output_path)}' Notepad mein kholo")
    print("  2. Sab text copy karo  (Ctrl+A -> Ctrl+C)")
    print("  3. InPage kholo -> Text box banao")
    print("  4. Paste karo  (Ctrl+V)")
    print("=" * 60)
    print()


if __name__ == '__main__':
    main()
