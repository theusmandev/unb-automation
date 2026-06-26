#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
==============================================================
  Unicode Urdu -> InPage Converter  |  Corrected Version
==============================================================
  Supports: .txt  aur  .docx  input files
  Output:   InPage-compatible .txt file (paste-ready)

  Install:  pip install python-docx    (sirf docx ke liye)

  YE VERSION KYU AYI:
  Purani script mein har Urdu letter ko ek single byte
  (e.g. 'ا': '\\xc7') diya gaya tha. Lekin asal InPage
  encoding aisi nahi -- InPage har character ko DO units
  mein store karta hai: control marker 0x04 + ek specific
  code. Isi liye purana mapping InPage mein paste karne par
  random/garbled symbols dikhata tha.

  Ye corrected mapping table verified open-source project se
  li gayi hai (UmerCodez/unicode-inpage-converter, GPL-3.0)
  aur Python mein dobara likhi gayi hai.
==============================================================
"""

import sys
import os

CTRL = '\u0004'  # InPage ka control marker -- har glyph code se pehle aata hai

# ─────────────────────────────────────────────────────────────────────────────
#  Multi-character sequences jo single-char loop se PEHLE handle hone chahiye
#  (warna unki individual parts alag se wrongly convert ho jayengi)
# ─────────────────────────────────────────────────────────────────────────────
PRE_REPLACEMENTS = [
    ('\u06CC\u0626', CTRL + '\u00A4' + CTRL + '\u00BF'),   # choti ye + hamza
    ('\u0623',        CTRL + '\u0081' + CTRL + '\u00BF'),   # alif + hamza above (أ)
    ('\u0624',        CTRL + '\u00A2' + CTRL + '\u00BF'),   # waw + hamza above (ؤ)
    ('\u0622',        CTRL + '\u0081' + CTRL + '\u00B3'),   # alif madda (آ)
]

# Characters consumed only inside PRE_REPLACEMENTS (used for unmapped-detection)
_PRE_REPLACEMENT_CHARS = {'\u06CC', '\u0626', '\u0623', '\u0624', '\u0622'}

# ─────────────────────────────────────────────────────────────────────────────
#  Verified Unicode -> InPage single-character mapping table
# ─────────────────────────────────────────────────────────────────────────────
UNICODE_TO_INPAGE = {
    # Noon ghunna, diacritics, Urdu digits
    '\u06BA': CTRL + '\u00A1', '\u064D': CTRL + '\u00A8',
    '\u06F1': CTRL + '\u00D1', '\u06F2': CTRL + '\u00D2', '\u06F3': CTRL + '\u00D3',
    '\u06F4': CTRL + '\u00D4', '\u06F5': CTRL + '\u00D5', '\u06F6': CTRL + '\u00D6',
    '\u06F7': CTRL + '\u00D7', '\u06F8': CTRL + '\u00D8', '\u06F9': CTRL + '\u00D9',
    '\u06F0': CTRL + '\u00D0',

    # Main alphabet
    '\u0642': CTRL + '\u203A',  # ق
    '\u0648': CTRL + '\u00A2',  # و
    '\u0639': CTRL + '\u02DC',  # ع
    '\u0631': CTRL + '\u017D',  # ر
    '\u062A': CTRL + '\u201E',  # ت
    '\u06D2': CTRL + '\u00A5',  # ے
    '\u0621': CTRL + '\u00A3',  # ء
    '\u06CC': CTRL + '\u00A4',  # ی
    '\u06C1': CTRL + '\u00A6',  # ہ
    '\u067E': CTRL + '\u0192',  # پ
    '\u060E': CTRL + '\u00F2',
    '\u0627': CTRL + '\u0081',  # ا
    '\u0633': CTRL + '\u2019',  # س
    '\u062F': CTRL + '\u2039',  # د
    '\u0641': CTRL + '\u0161',  # ف
    '\u06AF': CTRL + '\u009D',  # گ
    '\u06BE': CTRL + '\u00A7',  # ھ
    '\u062C': CTRL + '\u2021',  # ج
    '\u06A9': CTRL + '\u0153',  # ک
    '\u0644': CTRL + '\u017E',  # ل
    '\u061B': CTRL + '\u00EA',  # ؛
    '\u0632': CTRL + '\u0090',  # ز
    '\u0634': CTRL + '\u201C',  # ش
    '\u0686': CTRL + '\u02C6',  # چ
    '\u0637': CTRL + '\u2013',  # ط
    '\u0628': CTRL + '\u201A',  # ب
    '\u0646': CTRL + '\u00A0',  # ن
    '\u0645': CTRL + '\u0178',  # م
    '\u060C': CTRL + '\u00ED',  # ،
    '\u06D4': CTRL + '\u00F3',  # ۔
    '\u064B': CTRL + '\u00C7',
    '\u064A': CTRL + '\u00B8',  # ي (Arabic yeh)
    '\u0610': CTRL + '\u00F8',
    '\u0626': CTRL + '\u00A3',  # ئ
    '\u064C': CTRL + '\u00B5',
    '\u0651': CTRL + '\u00AD',  # shadda
    '\u0652': CTRL + '\u00B1',  # sukun
    '\uFDFA': CTRL + '\u00F6',
    '\u0611': CTRL + '\u00AE',
    '\u0691': CTRL + '\u008F',  # ڑ
    '\u0679': CTRL + '\u2026',  # ٹ
    '\u0601': CTRL + '\u00F7',
    '\u0657': CTRL + '\u00BE',
    '\u0670': CTRL + '\u00BD',
    '\u06C3': CTRL + '\u00B9',
    '\u064F': CTRL + '\u00AC',  # pesh
    '\u2018': CTRL + '\u00FD',  # left single quote
    '\u2019': CTRL + '\u00FE',  # right single quote
    '\u0614': CTRL + '\u00CF',
    '\u0653': CTRL + '\u00B3',
    '\u0635': CTRL + '\u201D',  # ص
    '\u0688': CTRL + '\u0152',  # ڈ
    '\u0656': CTRL + '\u00B0',
    '\u063A': CTRL + '\u2122',  # غ
    '\u062D': CTRL + '\u2030',  # ح
    '\u0636': CTRL + '\u2022',  # ض
    '\u062E': CTRL + '\u0160',  # خ
    '\u0612': CTRL + '\u00E7',
    '\u0630': CTRL + '\u008D',  # ذ
    '\u0698': CTRL + '\u2018',  # ژ
    '\u062B': CTRL + '\u2020',  # ث
    '\u0638': CTRL + '\u2014',  # ظ
    '\u0613': CTRL + '\u00E6',
    '\u0650': CTRL + '\u00AA',  # zer
    '\u064E': CTRL + '\u00AB',  # zabar
    '\u061F': CTRL + '\u00EE',  # ؟

    # Basic punctuation re-mapped to InPage codes
    '[': CTRL + '\u00FA', ']': CTRL + '\u00FB', '.': CTRL + '\u00FC',
    '!': CTRL + '\u00DA', ',': CTRL + '\u00F9', '/': CTRL + '\u00F1',
    ')': CTRL + '\u00E1', '(': CTRL + '\u00E2', ':': CTRL + '\u00E9',

    # Space also gets the control marker
    ' ': CTRL + ' ',
}

_RECOGNIZED_CHARS = set(UNICODE_TO_INPAGE.keys()) | _PRE_REPLACEMENT_CHARS


# ─────────────────────────────────────────────────────────────────────────────
#  Core Conversion
# ─────────────────────────────────────────────────────────────────────────────
def convert_unicode_to_inpage(text: str) -> tuple:
    # Detect truly unmapped characters BEFORE any replacement happens,
    # so control bytes produced by our own substitutions never get
    # mis-flagged as "unmapped".
    unmapped = set()
    for ch in text:
        if ch in ('\n', '\t', '\r'):
            continue
        if ch.strip() and ch not in _RECOGNIZED_CHARS:
            unmapped.add(ch)

    for old, new in PRE_REPLACEMENTS:
        text = text.replace(old, new)

    result = []
    for ch in text:
        result.append(UNICODE_TO_INPAGE.get(ch, ch))

    return ''.join(result), unmapped


# ─────────────────────────────────────────────────────────────────────────────
#  File Readers
# ─────────────────────────────────────────────────────────────────────────────
def read_txt_file(path: str) -> str:
    for enc in ['utf-8-sig', 'utf-8', 'utf-16', 'cp1256', 'iso-8859-6']:
        try:
            with open(path, 'r', encoding=enc) as f:
                content = f.read()
            print(f"    Encoding detect hua: {enc}")
            return content
        except (UnicodeDecodeError, LookupError):
            continue
    raise ValueError(f"File '{path}' ko kisi bhi encoding se nahi parh saka.")


def read_docx_file(path: str) -> str:
    try:
        from docx import Document
    except ImportError:
        print("\n  X python-docx install karo:")
        print("     pip install python-docx\n")
        sys.exit(1)

    doc = Document(path)
    lines = []
    for para in doc.paragraphs:
        lines.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                lines.append(row_text)

    return '\n'.join(lines)


# ─────────────────────────────────────────────────────────────────────────────
#  Main Logic
# ─────────────────────────────────────────────────────────────────────────────
def main():
    # 👇 Apne paths yahan set karein 👇
    input_path = r"C:\Users\PCS\Downloads\unicode_to_inpage converter\orignal.txt"
    output_path = r"C:\Users\PCS\Downloads\unicode_to_inpage converter\orignal_inpage.txt"

    # ── Checks ───────────────────────────────────────────────────────────────
    if not os.path.exists(input_path):
        print(f"\nX File nahi mili: '{input_path}'")
        print("! Barae meharbani check karein ke name aur location bilkul theek hai.")
        sys.exit(1)

    ext = os.path.splitext(input_path)[1].lower()
    if ext not in ('.txt', '.docx'):
        print("\nX Sirf .txt aur .docx files support hain.")
        sys.exit(1)

    print()
    print("=" * 58)
    print("   Unicode Urdu -> InPage Converter (Corrected)")
    print("=" * 58)
    print(f"  Input  : {input_path}")
    print(f"  Output : {output_path}")
    print("-" * 58)

    print("\n  File parh raha hoon...")
    try:
        if ext == '.txt':
            text = read_txt_file(input_path)
        else:
            print("    DOCX paragraphs extract kar raha hoon...")
            text = read_docx_file(input_path)
    except Exception as e:
        print(f"\n  X Error: {e}")
        sys.exit(1)

    word_count = len(text.split())
    char_count = len(text)
    print(f"    {char_count:,} characters | ~{word_count:,} words mile")

    print("\n  Convert ho raha hai...")
    converted, unmapped = convert_unicode_to_inpage(text)

    pct = round(100 - (len(unmapped) / max(char_count, 1) * 100), 1)
    print(f"    {pct}% characters successfully convert hue")

    if unmapped:
        sample = ', '.join(f"'{c}'(U+{ord(c):04X})" for c in sorted(unmapped)[:8])
        print(f"    Unmapped (as-is rakhe gaye): {sample}")

    # IMPORTANT: utf-8 NAHI use karna -- Notepad kabhi kabhi UTF-8 BOM ko
    # theek se detect nahi karta aur file ko ANSI samajh kar khol deta hai,
    # jis se UTF-8 ke multi-byte sequences "â€™" jaisi mojibake mein badal
    # jate hain jab InPage mein paste ki jaye. UTF-16 (with BOM) Notepad
    # mein hamesha reliably detect hoti hai, is liye yahi safe choice hai.
    print("\n  File save ho rahi hai...")
    try:
        with open(output_path, 'w', encoding='utf-16', newline='') as f:
            f.write(converted)
    except Exception as e:
        print(f"\n  X Save error: {e}")
        sys.exit(1)

    size_kb = os.path.getsize(output_path) / 1024
    print(f"    '{os.path.basename(output_path)}' save ho gayi! ({size_kb:.1f} KB)")

    print()
    print("=" * 58)
    print("  KAAM MUKAMMAL HO GAYA!")
    print("-" * 58)
    print("  InPage mein use karne ka tarika:")
    print(f"  1. '{os.path.basename(output_path)}' Notepad mein kholo")
    print("  2. Sab text copy karo  (Ctrl+A -> Ctrl+C)")
    print("  3. InPage kholo -> Text box banao")
    print("  4. Paste karo  (Ctrl+V)")
    print("=" * 58)
    print()


if __name__ == '__main__':
    main()
