#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
==============================================================
  Unicode Urdu → InPage Converter  |  Professional Tool
==============================================================
  Supports: .txt  aur  .docx  input files
  Output:   InPage-compatible .txt file (paste-ready)

  Install:  pip install python-docx    (sirf docx ke liye)

  Usage:
    python unicode_to_inpage.py input.txt
    python unicode_to_inpage.py input.docx
    python unicode_to_inpage.py input.txt --output result.txt
==============================================================
"""

import sys
import os
import argparse

# ─────────────────────────────────────────────────────────────────────────────
#  Unicode → InPage Encoding Table
#  (InPage uses its own legacy single-byte encoding for Urdu/Nastaliq)
# ─────────────────────────────────────────────────────────────────────────────
UNICODE_TO_INPAGE = {
    # Whitespace
    ' ':  ' ',
    '\t': '\t',
    '\n': '\n',
    '\r': '',

    # ── Hamza & Alef forms ───────────────────────────────────────────────────
    'ء': '\xc1',   # hamza
    'آ': '\xc2',   # alef madda
    'أ': '\xc3',   # alef hamza above
    'ؤ': '\xc4',   # waw hamza
    'إ': '\xc5',   # alef hamza below
    'ئ': '\xc6',   # ya hamza
    'ا': '\xc7',   # alef

    # ── Ba group ─────────────────────────────────────────────────────────────
    'ب': '\xc8',   # ba
    'پ': '\xc9',   # pa
    'ت': '\xca',   # ta
    'ٹ': '\xcb',   # tta
    'ث': '\xcc',   # sa (tse)

    # ── Jim group ────────────────────────────────────────────────────────────
    'ج': '\xcd',   # jim
    'چ': '\xce',   # che
    'ح': '\xcf',   # ha (huti)
    'خ': '\xd0',   # kha

    # ── Dal group ────────────────────────────────────────────────────────────
    'د': '\xd1',   # dal
    'ڈ': '\xd2',   # dda
    'ذ': '\xd3',   # zal

    # ── Ra group ─────────────────────────────────────────────────────────────
    'ر': '\xd4',   # ra
    'ڑ': '\xd5',   # rra
    'ز': '\xd6',   # zain
    'ژ': '\xd7',   # zha

    # ── Sin group ────────────────────────────────────────────────────────────
    'س': '\xd8',   # sin
    'ش': '\xd9',   # shin

    # ── Sad group ────────────────────────────────────────────────────────────
    'ص': '\xda',   # sad
    'ض': '\xdb',   # dad

    # ── Emphatic group ───────────────────────────────────────────────────────
    'ط': '\xdc',   # toa
    'ظ': '\xdd',   # zoa
    'ع': '\xde',   # ain
    'غ': '\xdf',   # ghain

    # ── Fa group ─────────────────────────────────────────────────────────────
    'ف': '\xe1',   # fa
    'ق': '\xe2',   # qaf

    # ── Kaf group ────────────────────────────────────────────────────────────
    'ک': '\xe3',   # kaf (Urdu)
    'ك': '\xe3',   # kaf (Arabic)
    'گ': '\xe4',   # ga

    # ── Lam & Mim ────────────────────────────────────────────────────────────
    'ل': '\xe5',   # lam
    'م': '\xe6',   # mim

    # ── Nun ──────────────────────────────────────────────────────────────────
    'ن': '\xe7',   # nun
    'ں': '\xe8',   # nun ghunna

    # ── Waw & Ha ─────────────────────────────────────────────────────────────
    'و': '\xe9',   # waw
    'ہ': '\xea',   # ha (Urdu gol he)
    'ه': '\xea',   # ha (Arabic)
    'ھ': '\xeb',   # do chashmi ha

    # ── Ya ───────────────────────────────────────────────────────────────────
    'ی': '\xed',   # ya (choti ye)
    'ي': '\xed',   # ya (Arabic)
    'ے': '\xee',   # barri ya

    # ── Diacritics (aeraab) ──────────────────────────────────────────────────
    'َ': '\xf0',   # zabar (fatha)
    'ِ': '\xf1',   # zer (kasra)
    'ُ': '\xf2',   # pesh (damma)
    'ً': '\xf3',   # tanwin zabar
    'ٍ': '\xf4',   # tanwin zer
    'ٌ': '\xf5',   # tanwin pesh
    'ّ': '\xf6',   # tashdid (shadda)
    'ْ': '\xf7',   # sukun

    # ── Urdu Punctuation ─────────────────────────────────────────────────────
    '۔': '\xae',   # Urdu full stop
    '،': '\xac',   # Urdu/Arabic comma
    '؟': '\xbf',   # Arabic question mark
    '؛': ';',
    '٪': '%',
    '٭': '*',
    '۞': '*',

    # ── Arabic-Indic numerals ────────────────────────────────────────────────
    '۰': '0', '۱': '1', '۲': '2', '۳': '3', '۴': '4',
    '۵': '5', '۶': '6', '۷': '7', '۸': '8', '۹': '9',

    # ── Western digits & basic punctuation (passthrough) ────────────────────
    **{str(d): str(d) for d in range(10)},
    '.': '.', ',': ',', '!': '!', '?': '?', ':': ':',
    ';': ';', '-': '-', '–': '-', '—': '-', '_': '_',
    '(': '(', ')': ')', '[': '[', ']': ']',
    '"': '"', "'": "'", '"': '"', '"': '"',
    ''': "'", ''': "'",
    '\u200c': '',   # zero-width non-joiner (remove)
    '\u200d': '',   # zero-width joiner (remove)
    '\u200e': '',   # left-to-right mark (remove)
    '\u200f': '',   # right-to-left mark (remove)
    '\ufeff': '',   # BOM (remove)
}

# 2-character ligatures (check before single chars)
LIGATURES = {
    'لا': '\xfe\xc7',
    'لآ': '\xfe\xc2',
    'لأ': '\xfe\xc3',
    'لإ': '\xfe\xc5',
}


# ─────────────────────────────────────────────────────────────────────────────
#  Core Conversion
# ─────────────────────────────────────────────────────────────────────────────
def convert_unicode_to_inpage(text: str) -> str:
    """Unicode Urdu string → InPage-compatible string."""
    result = []
    unmapped = set()
    i = 0

    while i < len(text):
        # Check 2-char ligatures first
        two = text[i:i+2]
        if two in LIGATURES:
            result.append(LIGATURES[two])
            i += 2
            continue

        ch = text[i]
        if ch in UNICODE_TO_INPAGE:
            result.append(UNICODE_TO_INPAGE[ch])
        else:
            result.append(ch)
            if ch.strip():
                unmapped.add(ch)
        i += 1

    return ''.join(result), unmapped


# ─────────────────────────────────────────────────────────────────────────────
#  File Readers
# ─────────────────────────────────────────────────────────────────────────────
def read_txt_file(path: str) -> str:
    """TXT file padhta hai — multiple encodings try karta hai."""
    for enc in ['utf-8-sig', 'utf-8', 'utf-16', 'cp1256', 'iso-8859-6']:
        try:
            with open(path, 'r', encoding=enc) as f:
                content = f.read()
            print(f"    Encoding detect hua: {enc}")
            return content
        except (UnicodeDecodeError, LookupError):
            continue
    raise ValueError(f"File '{path}' ko kisi bhi encoding se nahi parh saka.")


def read_docx_file(path: str) -> str:
    """DOCX file se text extract karta hai — paragraphs preserve karta hai."""
    try:
        from docx import Document
    except ImportError:
        print("\n  ❌ python-docx install karo:")
        print("     pip install python-docx\n")
        sys.exit(1)

    doc = Document(path)
    lines = []
    for para in doc.paragraphs:
        lines.append(para.text)   # blank lines bhi rakh do (structure preserve)

    # Tables ka text bhi extract karo
    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                lines.append(row_text)

    return '\n'.join(lines)


# ─────────────────────────────────────────────────────────────────────────────
#  Main
# ─────────────────────────────────────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser(
        description='Unicode Urdu ko InPage encoding mein convert karo',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Misal:
  python unicode_to_inpage.py article.txt
  python unicode_to_inpage.py novel.docx
  python unicode_to_inpage.py input.txt --output result.txt
        """
    )
    parser.add_argument('input',         help='Input file (.txt ya .docx)')
    parser.add_argument('--output', '-o',help='Output file naam (optional)')
    args = parser.parse_args()

    input_path = args.input

    # ── Checks ───────────────────────────────────────────────────────────────
    if not os.path.exists(input_path):
        print(f"\n❌ File nahi mili: '{input_path}'")
        sys.exit(1)

    ext = os.path.splitext(input_path)[1].lower()
    if ext not in ('.txt', '.docx'):
        print(f"\n❌ Sirf .txt aur .docx files support hain. Aapne diya: '{ext}'")
        sys.exit(1)

    # ── Output path ──────────────────────────────────────────────────────────
    if args.output:
        output_path = args.output
    else:
        base = os.path.splitext(input_path)[0]
        output_path = f"{base}_inpage.txt"

    # ── Header ───────────────────────────────────────────────────────────────
    print()
    print("═" * 58)
    print("   Unicode Urdu → InPage Converter")
    print("═" * 58)
    print(f"  📂 Input  : {input_path}")
    print(f"  💾 Output : {output_path}")
    print("─" * 58)

    # ── Read ─────────────────────────────────────────────────────────────────
    print("\n  📖 File parh raha hoon...")
    try:
        if ext == '.txt':
            text = read_txt_file(input_path)
        else:
            print("    DOCX paragraphs extract kar raha hoon...")
            text = read_docx_file(input_path)
    except Exception as e:
        print(f"\n  ❌ Error: {e}")
        sys.exit(1)

    word_count  = len(text.split())
    char_count  = len(text)
    print(f"    {char_count:,} characters | ~{word_count:,} words mile")

    # ── Convert ──────────────────────────────────────────────────────────────
    print("\n  🔄 Convert ho raha hai...")
    converted, unmapped = convert_unicode_to_inpage(text)

    mapped_count = sum(1 for ch in text if ch in UNICODE_TO_INPAGE or
                       text[max(0, text.index(ch)-1):text.index(ch)+1] in LIGATURES)
    pct = round(100 - (len(unmapped) / max(char_count, 1) * 100), 1)
    print(f"    ✅ {pct}% characters successfully convert hue")

    if unmapped:
        sample = ', '.join(f"'{c}'(U+{ord(c):04X})" for c in sorted(unmapped)[:8])
        print(f"    ℹ️  Unmapped (as-is rakhe gaye): {sample}")

    # ── Save ─────────────────────────────────────────────────────────────────
    print(f"\n  💾 File save ho rahi hai...")
    try:
        # latin-1 encoding use karo — ye 0x00-0xFF byte values perfectly preserve karta hai
        # jo ke InPage ki apni encoding ke liye zaroori hai
        with open(output_path, 'w', encoding='latin-1', errors='replace') as f:
            f.write(converted)
    except Exception as e:
        print(f"\n  ❌ Save error: {e}")
        sys.exit(1)

    size_kb = os.path.getsize(output_path) / 1024
    print(f"    ✅ '{output_path}' save ho gayi! ({size_kb:.1f} KB)")

    # ── Instructions ─────────────────────────────────────────────────────────
    print()
    print("═" * 58)
    print("  🎉 KAAM MUKAMMAL HO GAYA!")
    print("─" * 58)
    print("  InPage mein use karne ka tarika:")
    print(f"  1. '{output_path}' Notepad mein kholo")
    print("  2. Sab text copy karo  (Ctrl+A → Ctrl+C)")
    print("  3. InPage kholo → Text box banao")
    print("  4. Paste karo  (Ctrl+V)")
    print("  5. Font: Nafees Nastaliq / Jameel Noori Nastaleeq")
    print("     ya apna pasandida InPage Urdu font set karo")
    print("═" * 58)
    print()


if __name__ == '__main__':
    main()
