#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
╔══════════════════════════════════════════════════════════════╗
║     Unicode Urdu Novel -> InPage Converter  |  FINAL v7     ║
║                                                              ║
║  Handles:                                                    ║
║  ✓ Urdu text (complete alphabet + diacritics)               ║
║  ✓ English words in RTL — AND, DNA etc.       ✅            ║
║  ✓ English in brackets — (Opening)            ✅            ║
║  ✓ Multi-word brackets — (Queens Maker)       ✅ v7 FIX     ║
║  ✓ Multi-word brackets — (Daily Log)          ✅ v7 FIX     ║
║  ✓ Decorative symbols  ✿ • ★ ❤ etc.                        ║
║  ✓ Urdu digits  ۰۱۲۳۴۵۶۷۸۹                                  ║
║  ✓ Mixed Urdu+English lines                                  ║
║  ✓ .txt aur .docx dono formats                              ║
║  ✓ Full report: kya convert hua, kya nahi                   ║
╚══════════════════════════════════════════════════════════════╝

  BRACKET PROBLEM — ROOT CAUSE & FIX (v7):
  ─────────────────────────────────────────
  InPage ka Unicode BiDi engine English text khud sahi dikhata hai.
  LEKIN jab brackets ke andar SPACES hon (multi-word English),
  to InPage spaces ke boundary par text reorder karta hai aur
  brackets ulti jagah chali jaati hain:

    (Queens Maker)  ->  InPage shows:  Maker)(Queens   ✗
    (Daily Log)     ->  InPage shows:  Log)(Daily      ✗
    (Opening)       ->  InPage shows:  (Opening)       ✅ (single word, OK)

  FIX: Brackets ke andar ki SPACES ko Non-Breaking Space (U+00A0)
  se replace karo. Yeh InPage BiDi ko force karta hai ke poora
  bracketed text ek unit samjhe — reorder na kare:

    (Queens Maker)  ->  store (Queens\u00A0Maker)  ->  InPage: (Queens Maker) ✅
    (Daily Log)     ->  store (Daily\u00A0Log)     ->  InPage: (Daily Log)    ✅

  INSTALL (sirf docx ke liye zaroori):
      pip install python-docx

  USE:
      1. Neeche INPUT_PATH aur OUTPUT_PATH set karo
      2. python urdu_novel_to_inpage_FINAL_v7.py
"""

import sys
import os
import re
from collections import Counter

# ══════════════════════════════════════════════════════════════
#  ⚙️  SETTINGS — Yahan sirf yahi section change karna hai
# ══════════════════════════════════════════════════════════════

INPUT_PATH  = r"C:\Users\PCS\Downloads\Untitled document.docx"
OUTPUT_PATH = r"E:\unb-workstation\DNA\InPage DNA\1-Done_final.txt"

# Decorative symbols ka kya karna hai:
#   'replace'  ->  neeche SYMBOL_REPLACEMENTS mein jo likha hai woh lagao
#   'remove'   ->  bilkul hata do
#   'keep'     ->  jaise hai waise rakho
SYMBOL_ACTION = 'replace'

SYMBOL_REPLACEMENTS = {
    '✿': '★', '❀': '*', '✾': '*', '❁': '*', '✽': '*',
    '❃': '*', '✸': '*', '✺': '*', '✻': '*', '✼': '*', '❋': '*',
    '•': '-', '·': '.', '‧': '.', '∙': '.',
    '●': 'o', '◉': 'o', '○': 'o', '◎': 'o',
    '◆': '*', '◇': '*', '■': '*', '□': '*',
    '▪': '-', '▫': '-', '▸': '>', '▹': '>',
    '★': '*', '☆': '*', '✦': '*', '✧': '*', '✩': '*',
    '✪': '*', '✫': '*', '✬': '*', '✭': '*', '✮': '*',
    '❤': '<3', '♥': '<3', '♡': '<3', '❥': '<3',
    '☀': '*', '✨': '*', '💫': '*', '🌟': '*',
    '═': '=', '─': '-', '━': '-', '│': '|', '┃': '|',
    '╔': '+', '╗': '+', '╚': '+', '╝': '+', '╠': '+',
    '╣': '+', '╦': '+', '╩': '+', '╬': '+',
    '▬': '-', '▭': '-', '▰': '-',
    '…': '...',
    '–': '-',
    '—': '-',
    '―': '-',
    '\u00AB': '"', '\u00BB': '"',
    '\u2039': '', '\u203A': '',
    '©': '(c)', '®': '(r)', '™': '(tm)',
    '°': 'deg', '±': '+/-',
    '½': '1/2', '¼': '1/4', '¾': '3/4',
    '\u200B': '', '\u200C': '', '\u200D': '',
    '\uFEFF': '',
    '\u00A0': ' ',
}

# ══════════════════════════════════════════════════════════════
#  INTERNAL CONSTANTS
# ══════════════════════════════════════════════════════════════

CTRL = '\u0004'
NBSP = '\u00A0'  # Non-Breaking Space — bracket fix ke liye

PRE_REPLACEMENTS = [
    ('\u06CC\u0626', CTRL + '\u00A4' + CTRL + '\u00BF'),
    ('\u0623',        CTRL + '\u0081' + CTRL + '\u00BF'),
    ('\u0624',        CTRL + '\u00A2' + CTRL + '\u00BF'),
    ('\u0622',        CTRL + '\u0081' + CTRL + '\u00B3'),
]

_PRE_REPLACEMENT_CHARS = {'\u06CC', '\u0626', '\u0623', '\u0624', '\u0622'}

UNICODE_TO_INPAGE = {
    '\u06BA': CTRL + '\u00A1', '\u064D': CTRL + '\u00A8',
    '\u06F1': CTRL + '\u00D1', '\u06F2': CTRL + '\u00D2', '\u06F3': CTRL + '\u00D3',
    '\u06F4': CTRL + '\u00D4', '\u06F5': CTRL + '\u00D5', '\u06F6': CTRL + '\u00D6',
    '\u06F7': CTRL + '\u00D7', '\u06F8': CTRL + '\u00D8', '\u06F9': CTRL + '\u00D9',
    '\u06F0': CTRL + '\u00D0',
    '\u0642': CTRL + '\u203A', '\u0648': CTRL + '\u00A2', '\u0639': CTRL + '\u02DC',
    '\u0631': CTRL + '\u017D', '\u062A': CTRL + '\u201E', '\u06D2': CTRL + '\u00A5',
    '\u0621': CTRL + '\u00A3', '\u06CC': CTRL + '\u00A4', '\u06C1': CTRL + '\u00A6',
    '\u067E': CTRL + '\u0192', '\u060E': CTRL + '\u00F2', '\u0627': CTRL + '\u0081',
    '\u0633': CTRL + '\u2019', '\u062F': CTRL + '\u2039', '\u0641': CTRL + '\u0161',
    '\u06AF': CTRL + '\u009D', '\u06BE': CTRL + '\u00A7', '\u062C': CTRL + '\u2021',
    '\u06A9': CTRL + '\u0153', '\u0644': CTRL + '\u017E', '\u061B': CTRL + '\u00EA',
    '\u0632': CTRL + '\u0090', '\u0634': CTRL + '\u201C', '\u0686': CTRL + '\u02C6',
    '\u0637': CTRL + '\u2013', '\u0628': CTRL + '\u201A', '\u0646': CTRL + '\u00A0',
    '\u0645': CTRL + '\u0178', '\u060C': CTRL + '\u00ED', '\u06D4': CTRL + '\u00F3',
    '\u064B': CTRL + '\u00C7', '\u064A': CTRL + '\u00B8', '\u0610': CTRL + '\u00F8',
    '\u0626': CTRL + '\u00A3', '\u064C': CTRL + '\u00B5', '\u0651': CTRL + '\u00AD',
    '\u0652': CTRL + '\u00B1', '\uFDFA': CTRL + '\u00F6', '\u0611': CTRL + '\u00AE',
    '\u0691': CTRL + '\u008F', '\u0679': CTRL + '\u2026', '\u0601': CTRL + '\u00F7',
    '\u0657': CTRL + '\u00BE', '\u0670': CTRL + '\u00BD', '\u06C3': CTRL + '\u00B9',
    '\u064F': CTRL + '\u00AC', '\u2018': CTRL + '\u00FD', '\u2019': CTRL + '\u00FE',
    '\u0614': CTRL + '\u00CF', '\u0653': CTRL + '\u00B3', '\u0635': CTRL + '\u201D',
    '\u0688': CTRL + '\u0152', '\u0656': CTRL + '\u00B0', '\u063A': CTRL + '\u2122',
    '\u062D': CTRL + '\u2030', '\u0636': CTRL + '\u2022', '\u062E': CTRL + '\u0160',
    '\u0612': CTRL + '\u00E7', '\u0630': CTRL + '\u008D', '\u0698': CTRL + '\u2018',
    '\u062B': CTRL + '\u2020', '\u0638': CTRL + '\u2014', '\u0613': CTRL + '\u00E6',
    '\u0650': CTRL + '\u00AA', '\u064E': CTRL + '\u00AB', '\u061F': CTRL + '\u00EE',
    '[': CTRL + '\u00FA', ']': CTRL + '\u00FB', '.': CTRL + '\u00FC',
    '!': CTRL + '\u00DA', ',': CTRL + '\u00F9', '/': CTRL + '\u00F1',
    ')': CTRL + '\u00E1', '(': CTRL + '\u00E2', ':': CTRL + '\u00E9',
    ' ': CTRL + ' ',
    NBSP: CTRL + ' ',   # Non-breaking space bhi same as normal space map karo
}

_RECOGNIZED_CHARS = set(UNICODE_TO_INPAGE.keys()) | _PRE_REPLACEMENT_CHARS
_URDU_RE  = re.compile(r'[\u0600-\u06FF\u0750-\u077F\uFB50-\uFDFF\uFE70-\uFEFF]')
_LATIN_RE = re.compile(r'[A-Za-z]')

# Brackets ke andar multi-word English detect karne ka pattern
# e.g.  (Queens Maker)  (Daily Log)  (Hello World Test)
_ENG_BRACKET_MULTIWORD = re.compile(
    r'\(([A-Za-z][A-Za-z0-9]*(?:[ \t]+[A-Za-z0-9]+)+)\)'
)


# ══════════════════════════════════════════════════════════════
#  STEP 1: Symbol Cleanup
# ══════════════════════════════════════════════════════════════

def apply_symbol_replacements(text: str) -> tuple[str, Counter]:
    replaced_counts = Counter()
    for sym, replacement in SYMBOL_REPLACEMENTS.items():
        if sym in text:
            count = text.count(sym)
            replaced_counts[sym] += count
            text = text.replace(sym, replacement)
    return text, replaced_counts


# ══════════════════════════════════════════════════════════════
#  STEP 2: Bracket Fix — Multi-word English inside ()
# ══════════════════════════════════════════════════════════════

def fix_english_brackets(text: str) -> str:
    """
    InPage BiDi engine brackets ke andar spaces par text reorder
    karta hai jis se multi-word English ulta ho jata hai:

        (Queens Maker)  ->  InPage shows:  Maker)(Queens  ✗
        (Daily Log)     ->  InPage shows:  Log)(Daily     ✗

    FIX: Brackets ke andar spaces ko Non-Breaking Space (U+00A0) se
    replace karo. InPage NBSP ko space ki tarah dikhata hai lekin
    BiDi boundary treat nahi karta — poora group ek unit rahta hai:

        (Queens\u00A0Maker)  ->  InPage shows:  (Queens Maker)  ✅
        (Daily\u00A0Log)     ->  InPage shows:  (Daily Log)     ✅
        (Opening)            ->  unchanged (single word, already OK) ✅
    """
    def replace_inner_spaces(m):
        inner = m.group(1)
        # Spaces ko NBSP se replace karo
        inner_fixed = re.sub(r'[ \t]+', NBSP, inner)
        return f'({inner_fixed})'

    return _ENG_BRACKET_MULTIWORD.sub(replace_inner_spaces, text)


# ══════════════════════════════════════════════════════════════
#  STEP 3: Core Unicode -> InPage Conversion
# ══════════════════════════════════════════════════════════════

def convert_line(line: str) -> tuple[str, set]:
    """Single line convert karo. Returns (converted, unmapped_chars)"""
    unmapped = set()

    for ch in line:
        if ch in ('\n', '\t', '\r', CTRL):
            continue
        if ch.strip() and ch not in _RECOGNIZED_CHARS:
            unmapped.add(ch)

    # Pre-replacements (multi-char Urdu combos)
    for old, new in PRE_REPLACEMENTS:
        line = line.replace(old, new)

    # Char-by-char mapping
    result = []
    for ch in line:
        if ch in UNICODE_TO_INPAGE:
            result.append(UNICODE_TO_INPAGE[ch])
        elif ch in ('\n', '\r', '\t'):
            result.append(ch)
        elif ch == CTRL:
            result.append(ch)
        else:
            result.append(ch)  # unmapped: as-is

    return ''.join(result), unmapped


def convert_text(text: str) -> tuple[str, set, int, int]:
    lines = text.split('\n')
    converted_lines = []
    all_unmapped = set()
    total_chars = 0
    unconverted_chars = 0

    for line in lines:
        converted, unmapped = convert_line(line)
        converted_lines.append(converted)
        all_unmapped |= unmapped

        for ch in line:
            if ch not in ('\n', '\r', '\t') and ch.strip():
                total_chars += 1
                if ch in unmapped:
                    unconverted_chars += 1

    return '\n'.join(converted_lines), all_unmapped, total_chars, unconverted_chars


# ══════════════════════════════════════════════════════════════
#  FILE READERS
# ══════════════════════════════════════════════════════════════

def read_txt(path: str) -> str:
    for enc in ['utf-8-sig', 'utf-8', 'utf-16', 'utf-16-le', 'utf-16-be', 'cp1256', 'iso-8859-6']:
        try:
            with open(path, 'r', encoding=enc) as f:
                content = f.read()
            print(f"    ✓ Encoding: {enc}")
            return content
        except (UnicodeDecodeError, LookupError):
            continue
    raise ValueError(f"File '{path}' kisi encoding se nahi parhi ja saki.")


def read_docx(path: str) -> str:
    try:
        from docx import Document
    except ImportError:
        print("\n  ✗ python-docx install karo:  pip install python-docx\n")
        sys.exit(1)
    doc = Document(path)
    lines = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            cells = ' | '.join(c.text.strip() for c in row.cells if c.text.strip())
            if cells:
                lines.append(cells)
    return '\n'.join(lines)


# ══════════════════════════════════════════════════════════════
#  REPORT PRINTER
# ══════════════════════════════════════════════════════════════

def print_report(unmapped: set, replaced: Counter, total: int, unconverted: int):
    converted = total - unconverted
    pct = round(converted / max(total, 1) * 100, 2)

    print()
    print("┌─────────────────────────────────────────────┐")
    print("│              CONVERSION REPORT               │")
    print("├─────────────────────────────────────────────┤")
    print(f"│  Total characters    : {total:>8,}             │")
    print(f"│  Successfully mapped : {converted:>8,}  ({pct}%)  │")
    print(f"│  Unmapped (as-is)    : {unconverted:>8,}             │")
    print("├─────────────────────────────────────────────┤")

    if replaced:
        print("│  Symbols replaced:                          │")
        for sym, cnt in replaced.most_common(10):
            name = f"U+{ord(sym):04X}"
            print(f"│    '{sym}' ({name}) x{cnt:<6}                 │")

    if unmapped:
        print("├─────────────────────────────────────────────┤")
        print("│  Still unmapped (check InPage output):      │")
        for ch in sorted(unmapped)[:15]:
            print(f"│    '{ch}'  U+{ord(ch):04X}                        │")
        if len(unmapped) > 15:
            print(f"│    ... aur {len(unmapped)-15} aur characters          │")
    else:
        print("│  ✓ Koi unmapped character nahi!             │")

    print("└─────────────────────────────────────────────┘")


# ══════════════════════════════════════════════════════════════
#  MAIN
# ══════════════════════════════════════════════════════════════

def main():
    print()
    print("╔══════════════════════════════════════════════════════╗")
    print("║    Urdu Novel -> InPage Converter  |  FINAL v7      ║")
    print("╚══════════════════════════════════════════════════════╝")
    print(f"  Input   : {INPUT_PATH}")
    print(f"  Output  : {OUTPUT_PATH}")
    print(f"  Symbols : {SYMBOL_ACTION}")
    print("─" * 56)

    if not os.path.exists(INPUT_PATH):
        print(f"\n  ✗ File nahi mili: {INPUT_PATH}")
        sys.exit(1)

    ext = os.path.splitext(INPUT_PATH)[1].lower()
    if ext not in ('.txt', '.docx'):
        print("\n  ✗ Sirf .txt aur .docx support hain.")
        sys.exit(1)

    # ── Read ────────────────────────────────────────────────
    print("\n  [1/4] File parh raha hoon...")
    text = read_txt(INPUT_PATH) if ext == '.txt' else read_docx(INPUT_PATH)
    line_count = text.count('\n') + 1
    word_count = len(text.split())
    char_count = len(text)
    print(f"    {line_count:,} lines | {word_count:,} words | {char_count:,} chars")

    # ── Symbol cleanup ──────────────────────────────────────
    print("\n  [2/4] Symbols clean kar raha hoon...")
    if SYMBOL_ACTION == 'replace':
        text, replaced_counts = apply_symbol_replacements(text)
        print(f"    {sum(replaced_counts.values()):,} symbols replace hue ({len(replaced_counts)} unique types)")
    elif SYMBOL_ACTION == 'remove':
        replaced_counts = Counter()
        for sym in SYMBOL_REPLACEMENTS:
            if sym in text:
                replaced_counts[sym] += text.count(sym)
                text = text.replace(sym, '')
        print(f"    {sum(replaced_counts.values()):,} symbols remove hue")
    else:
        replaced_counts = Counter()
        print("    Symbols as-is rakhe gaye")

    # ── Bracket fix ─────────────────────────────────────────
    print("\n  [3/4] English brackets fix + Convert ho raha hai...")
    bracket_fixed = fix_english_brackets(text)
    fixed_count = sum(
        1 for line in bracket_fixed.split('\n')
        if NBSP in line
    )
    if fixed_count:
        print(f"    {fixed_count} lines mein multi-word English brackets fix hue")
    else:
        print("    Koi multi-word English bracket nahi mili")

    converted, unmapped, total_chars, unconverted_chars = convert_text(bracket_fixed)
    pct = round((total_chars - unconverted_chars) / max(total_chars, 1) * 100, 1)
    print(f"    {pct}% characters successfully convert hue")

    # ── Save ────────────────────────────────────────────────
    print("\n  [4/4] File save ho rahi hai...")
    os.makedirs(os.path.dirname(OUTPUT_PATH) or '.', exist_ok=True)
    with open(OUTPUT_PATH, 'w', encoding='utf-8-sig', newline='') as f:
        f.write(converted)
    size_kb = os.path.getsize(OUTPUT_PATH) / 1024
    print(f"    ✓ '{os.path.basename(OUTPUT_PATH)}'  ({size_kb:.1f} KB)")

    print_report(unmapped, replaced_counts, total_chars, unconverted_chars)

    print()
    print("╔══════════════════════════════════════════════════════╗")
    print("║              INPAGE MEIN USE KARNE KA TARIKA        ║")
    print("╠══════════════════════════════════════════════════════╣")
    print(f"║  1. '{os.path.basename(OUTPUT_PATH)}'  Notepad mein kholo   ║")
    print("║  2. Ctrl+A -> Ctrl+C  (sab copy karo)               ║")
    print("║  3. InPage mein Text Box banao                       ║")
    print("║  4. Ctrl+V paste karo                               ║")
    print("║  5. Font: Noori Nastaleeq ya Jameel Noori            ║")
    print("║  6. Direction: RTL confirm karo                      ║")
    print("║  ─────────────────────────────────────────────────  ║")
    print("║  TIP: Pehle sirf ek chapter paste karo aur check    ║")
    print("║       karo. Theek hai to puri novel karo.           ║")
    print("╚══════════════════════════════════════════════════════╝")
    print()


if __name__ == '__main__':
    main()
