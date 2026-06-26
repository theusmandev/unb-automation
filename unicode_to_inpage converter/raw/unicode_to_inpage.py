#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
==============================================================
  Unicode Urdu → InPage Converter  |  Professional Tool
==============================================================
  Supports: .txt  aur  .docx  input files
  Output:   InPage-compatible .txt file (paste-ready)

  Install:  pip install python-docx    (sirf docx ke liye)
==============================================================
"""

import sys
import os

# ─────────────────────────────────────────────────────────────────────────────
#  Unicode → InPage Encoding Table
# ─────────────────────────────────────────────────────────────────────────────
UNICODE_TO_INPAGE = {
    # Whitespace
    ' ':  ' ', '\t': '\t', '\n': '\n', '\r': '',

    # ── Hamza & Alef forms ───────────────────────────────────────────────────
    'ء': '\xc1', 'آ': '\xc2', 'أ': '\xc3', 'ؤ': '\xc4', 'إ': '\xc5', 'ئ': '\xc6', 'ا': '\xc7',

    # ── Ba group ─────────────────────────────────────────────────────────────
    'ب': '\xc8', 'پ': '\xc9', 'ت': '\xca', 'ٹ': '\xcb', 'ث': '\xcc',

    # ── Jim group ────────────────────────────────────────────────────────────
    'ج': '\xcd', 'چ': '\xce', 'ح': '\xcf', 'خ': '\xd0',

    # ── Dal group ────────────────────────────────────────────────────────────
    'د': '\xd1', 'ڈ': '\xd2', 'ذ': '\xd3',

    # ── Ra group ─────────────────────────────────────────────────────────────
    'ر': '\xd4', 'ڑ': '\xd5', 'ز': '\xd6', 'ژ': '\xd7',

    # ── Sin group ────────────────────────────────────────────────────────────
    'س': '\xd8', 'ش': '\xd9',

    # ── Sad group ────────────────────────────────────────────────────────────
    'ص': '\xda', 'ض': '\xdb',

    # ── Emphatic group ───────────────────────────────────────────────────────
    'ط': '\xdc', 'ظ': '\xdd', 'ع': '\xde', 'غ': '\xdf',

    # ── Fa group ─────────────────────────────────────────────────────────────
    'ف': '\xe1', 'ق': '\xe2',

    # ── Kaf group ────────────────────────────────────────────────────────────
    'ک': '\xe3', 'ك': '\xe3', 'گ': '\xe4',

    # ── Lam & Mim ────────────────────────────────────────────────────────────
    'ل': '\xe5', 'م': '\xe6',

    # ── Nun ──────────────────────────────────────────────────────────────────
    'ن': '\xe7', 'ں': '\xe8',

    # ── Waw & Ha ─────────────────────────────────────────────────────────────
    'و': '\xe9', 'ہ': '\xea', 'ه': '\xea', 'ھ': '\xeb',

    # ── Ya ───────────────────────────────────────────────────────────────────
    'ی': '\xed', 'ي': '\xed', 'ے': '\xee',

    # ── Diacritics (aeraab) ──────────────────────────────────────────────────
    'َ': '\xf0', 'ِ': '\xf1', 'ُ': '\xf2', 'ً': '\xf3', 'ٍ': '\xf4', 'ٌ': '\xf5', 'ّ': '\xf6', 'ْ': '\xf7',

    # ── Urdu Punctuation ─────────────────────────────────────────────────────
    '۔': '\xae', '،': '\xac', '؟': '\xbf', '؛': ';', '٪': '%', '٭': '*', '۞': '*',

    # ── Arabic-Indic numerals ────────────────────────────────────────────────
    '۰': '0', '۱': '1', '۲': '2', '۳': '3', '۴': '4',
    '۵': '5', '۶': '6', '۷': '7', '۸': '8', '۹': '9',

    # ── Western digits & basic punctuation (passthrough) ────────────────────
    **{str(d): str(d) for d in range(10)},
    '.': '.', ',': ',', '!': '!', '?': '?', ':': ':',
    ';': ';', '-': '-', '–': '-', '—': '-', '_': '_',
    '(': '(', ')': ')', '[': '[', ']': ']',
    '"': '"', "'": "'", '“': '"', '”': '"',
    '‘': "'", '’': "'",
    '\u200c': '', '\u200d': '', '\u200e': '', '\u200f': '', '\ufeff': '',
}

LIGATURES = {
    'لا': '\xfe\xc7', 'لآ': '\xfe\xc2', 'لأ': '\xfe\xc3', 'لإ': '\xfe\xc5',
}


# ─────────────────────────────────────────────────────────────────────────────
#  Core Conversion
# ─────────────────────────────────────────────────────────────────────────────
def convert_unicode_to_inpage(text: str) -> tuple:
    result = []
    unmapped = set()
    i = 0

    while i < len(text):
        two = text[i:i+2]
        if two in LIGATURES:
            result.append(LIGATURES[two])
            i += 2
            continue

        ch = text[i]
        if ch in UNICODE_TO_INPAGE:
            result.append(UNICODE_TO_INPAGE[ch])
        else:
            result.append(ch)
            if ch.strip():
                unmapped.add(ch)
        i += 1

    return ''.join(result), unmapped


# ─────────────────────────────────────────────────────────────────────────────
#  File Readers
# ─────────────────────────────────────────────────────────────────────────────
def read_txt_file(path: str) -> str:
    for enc in ['utf-8-sig', 'utf-8', 'utf-16', 'cp1256', 'iso-8859-6']:
        try:
            with open(path, 'r', encoding=enc) as f:
                content = f.read()
            print(f"    Encoding detect hua: {enc}")
            return content
        except (UnicodeDecodeError, LookupError):
            continue
    raise ValueError(f"File '{path}' ko kisi bhi encoding se nahi parh saka.")

def read_docx_file(path: str) -> str:
    try:
        from docx import Document
    except ImportError:
        print("\n  ❌ python-docx install karo:")
        print("     pip install python-docx\n")
        sys.exit(1)

    doc = Document(path)
    lines = []
    for para in doc.paragraphs:
        lines.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                lines.append(row_text)

    return '\n'.join(lines)


# ─────────────────────────────────────────────────────────────────────────────
#  Main Logic (HARDCODED PATHS)
# ─────────────────────────────────────────────────────────────────────────────
def main():
    # 👇 YAHAN AAPKE HARDCODED PATHS HAIN 👇
    
    # Input file ka path ('r' lagana zaroori hai taake \U error na aaye)
    input_path = r"C:\Users\PCS\Downloads\unicode_to_inpage converter\orignal.txt"
    
    # Output file ka path usi folder mein rakha hai
    output_path = r"C:\Users\PCS\Downloads\unicode_to_inpage converter\orignal_inpage.txt"

    # ─────────────────────────────────────────────────────────────────────────

    # ── Checks ───────────────────────────────────────────────────────────────
    if not os.path.exists(input_path):
        print(f"\n❌ File nahi mili: '{input_path}'")
        print("❗ Barae meharbani check karein ke name aur location bilkul theek hai.")
        sys.exit(1)

    ext = os.path.splitext(input_path)[1].lower()
    if ext not in ('.txt', '.docx'):
        print(f"\n❌ Sirf .txt aur .docx files support hain.")
        sys.exit(1)

    # ── Header ───────────────────────────────────────────────────────────────
    print()
    print("═" * 58)
    print("   Unicode Urdu → InPage Converter")
    print("═" * 58)
    print(f"  📂 Input  : {input_path}")
    print(f"  💾 Output : {output_path}")
    print("─" * 58)

    # ── Read ─────────────────────────────────────────────────────────────────
    print("\n  📖 File parh raha hoon...")
    try:
        if ext == '.txt':
            text = read_txt_file(input_path)
        else:
            print("    DOCX paragraphs extract kar raha hoon...")
            text = read_docx_file(input_path)
    except Exception as e:
        print(f"\n  ❌ Error: {e}")
        sys.exit(1)

    word_count  = len(text.split())
    char_count  = len(text)
    print(f"    {char_count:,} characters | ~{word_count:,} words mile")

    # ── Convert ──────────────────────────────────────────────────────────────
    print("\n  🔄 Convert ho raha hai...")
    converted, unmapped = convert_unicode_to_inpage(text)

    pct = round(100 - (len(unmapped) / max(char_count, 1) * 100), 1)
    print(f"    ✅ {pct}% characters successfully convert hue")

    if unmapped:
        sample = ', '.join(f"'{c}'(U+{ord(c):04X})" for c in sorted(unmapped)[:8])
        print(f"    ℹ️  Unmapped (as-is rakhe gaye): {sample}")

    # ── Save ─────────────────────────────────────────────────────────────────
    print(f"\n  💾 File save ho rahi hai...")
    try:
        with open(output_path, 'w', encoding='latin-1', errors='replace') as f:
            f.write(converted)
    except Exception as e:
        print(f"\n  ❌ Save error: {e}")
        sys.exit(1)

    size_kb = os.path.getsize(output_path) / 1024
    print(f"    ✅ '{os.path.basename(output_path)}' save ho gayi! ({size_kb:.1f} KB)")

    # ── Instructions ─────────────────────────────────────────────────────────
    print()
    print("═" * 58)
    print("  🎉 KAAM MUKAMMAL HO GAYA!")
    print("─" * 58)
    print("  InPage mein use karne ka tarika:")
    print(f"  1. '{os.path.basename(output_path)}' Notepad mein kholo")
    print("  2. Sab text copy karo  (Ctrl+A → Ctrl+C)")
    print("  3. InPage kholo → Text box banao")
    print("  4. Paste karo  (Ctrl+V)")
    print("═" * 58)
    print()


if __name__ == '__main__':
    main()