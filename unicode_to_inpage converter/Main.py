#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
╔══════════════════════════════════════════════════════════════╗
║     Unicode Urdu Novel -> InPage Converter  |  v4.3          ║
║                                                              ║
║  Handles:                                                    ║
║  ✓ Urdu text (complete alphabet + diacritics)                ║
║  ✓ Arabic text in duas (ه ك ة ى إ etc.)                      ║
║  ✓ English text RTL fix (Congrats Alia! sahi dikhega)        ║
║  ✓ Missing Space Fix: بے لوث ab بیلوث nahi banega            ║
║  ✓ Typo Auto-Fixer: OCR/Typing errors fix dictionary (v4.3)  ║
╚══════════════════════════════════════════════════════════════╝

  INSTALL (sirf docx ke liye zaroori):
      pip install python-docx


"""

import sys
import os
import re
from collections import Counter

# ══════════════════════════════════════════════════════════════
#  ⚙️  SETTINGS — Yahan sirf yahi section change karna hai
# ══════════════════════════════════════════════════════════════

INPUT_PATH  = r"C:\Users\PCS\Downloads\khird-e-junoon-by-nayaab-saeed-Episode 3.docx"
OUTPUT_PATH = r"C:\Users\PCS\Downloads\khird-e-junoon-by-nayaab-saeed-Episode 3.txt"
# V4.3 FIX: Agar Word file mein OCR ya typist ki ghaltiyan hain, unhe yahan theek karein
TYPO_FIXES = {
    'کلاشٹ': 'کلازٹ',
    'کلاسٹ': 'کلازٹ',
    # Aap yahan aur bhi words add kar sakte hain, misaal ke taur par:
    # 'ghalat_word': 'sahi_word',
}

# Decorative symbols (✿ • ★ etc.) ka kya karna hai:
SYMBOL_ACTION = 'replace'

SYMBOL_REPLACEMENTS = {
    # Decorative / floral
    '✿': '☆', '❀': '*', '✾': '*', '❁': '*', '✽': '*',
    '❃': '*', '✸': '*', '✺': '*', '✻': '*', '✼': '*',
    '❋': '*',
    # Bullets / dots
    '•': '-', '·': '.', '‧': '.', '∙': '.',
    '●': 'o', '◉': 'o', '○': 'o', '◎': 'o',
    '◆': '*', '◇': '*', '■': '*', '□': '*',
    '▪': '-', '▫': '-', '▸': '>', '▹': '>',
    # Stars / hearts
    '★': '*', '☆': '*', '✦': '*', '✧': '*', '✩': '*',
    '✪': '*', '✫': '*', '✬': '*', '✭': '*', '✮': '*',
    '❤': '<3', '♥': '<3', '♡': '<3', '❥': '<3',
    '☀': '*', '✨': '*', '💫': '*', '🌟': '*',
    # Box-drawing / separators
    '═': '=', '─': '-', '━': '-', '│': '|', '┃': '|',
    '╔': '+', '╗': '+', '╚': '+', '╝': '+', '╠': '+',
    '╣': '+', '╦': '+', '╩': '+', '╬': '+',
    '▬': '-', '▭': '-', '▰': '-',
    # Typography
    '…': '...',   # ellipsis -> 3 dots
    '–': '-',     # en dash
    '—': '-',     # em dash
    '―': '-',     # horizontal bar
    '\u00AB': '"', '\u00BB': '"',  # « »
    '\u2039': '',  # ‹
    '\u203A': '',  # ›
    # Misc
    '©': '(c)', '®': '(r)', '™': '(tm)',
    '°': 'deg', '±': '+/-',
    '½': '1/2', '¼': '1/4', '¾': '3/4',
    
    # --- Invisible characters ko delete karne ki bajaye space se replace ---
    '\u200B': ' ',  # zero-width space -> space
    '\u200C': ' ',  # zero-width non-joiner -> space
    '\u200D': '',   # zero-width joiner
    '\uFEFF': '',   # BOM
    '\u00A0': ' ',  # non-breaking space
}

# ══════════════════════════════════════════════════════════════
#  INTERNAL CONSTANTS
# ══════════════════════════════════════════════════════════════

CTRL = '\u0004'

PRE_REPLACEMENTS = [
    ('\u06CC\u0626', CTRL + '\u00A4' + CTRL + '\u00BF'),
    ('\u0623',        CTRL + '\u0081' + CTRL + '\u00BF'),  # أ Alef hamza above
    ('\u0624',        CTRL + '\u00A2' + CTRL + '\u00BF'),  # ؤ Waw hamza
    ('\u0622',        CTRL + '\u0081' + CTRL + '\u00B3'),  # آ Alef madda
]

_PRE_REPLACEMENT_CHARS = {'\u06CC', '\u0626', '\u0623', '\u0624', '\u0622'}

UNICODE_TO_INPAGE = {
    # ── Urdu-specific letters ──────────────────────────────
    '\u06BA': CTRL + '\u00A1',  # ں  Noon Ghunna
    '\u064D': CTRL + '\u00A8',  # ٍ  Kasratan
    '\u06F1': CTRL + '\u00D1',  # ۱
    '\u06F2': CTRL + '\u00D2',  # ۲
    '\u06F3': CTRL + '\u00D3',  # ۳
    '\u06F4': CTRL + '\u00D4',  # ۴
    '\u06F5': CTRL + '\u00D5',  # ۵
    '\u06F6': CTRL + '\u00D6',  # ۶
    '\u06F7': CTRL + '\u00D7',  # ۷
    '\u06F8': CTRL + '\u00D8',  # ۸
    '\u06F9': CTRL + '\u00D9',  # ۹
    '\u06F0': CTRL + '\u00D0',  # ۰
    '\u0642': CTRL + '\u203A',  # ق
    '\u0648': CTRL + '\u00A2',  # و
    '\u0639': CTRL + '\u02DC',  # ع
    '\u0631': CTRL + '\u017D',  # ر
    '\u062A': CTRL + '\u201E',  # ت
    '\u06D2': CTRL + '\u00A5',  # ے
    '\u0621': CTRL + '\u00A3',  # ء
    '\u06CC': CTRL + '\u00A4',  # ی  (Urdu Ye)
    '\u06C1': CTRL + '\u00A6',  # ہ  (Urdu He)
    '\u067E': CTRL + '\u0192',  # پ
    '\u060E': CTRL + '\u00F2',  # ؎
    '\u0627': CTRL + '\u0081',  # ا
    '\u0633': CTRL + '\u2019',  # س
    '\u062F': CTRL + '\u2039',  # د
    '\u0641': CTRL + '\u0161',  # ف
    '\u06AF': CTRL + '\u009D',  # گ
    '\u06BE': CTRL + '\u00A7',  # ھ  (Do Chashmi He)
    '\u062C': CTRL + '\u2021',  # ج
    '\u06A9': CTRL + '\u0153',  # ک  (Urdu Kaf)
    '\u0644': CTRL + '\u017E',  # ل
    '\u061B': CTRL + '\u00EA',  # ؛
    '\u0632': CTRL + '\u0090',  # ز
    '\u0634': CTRL + '\u201C',  # ش
    '\u0686': CTRL + '\u02C6',  # چ
    '\u0637': CTRL + '\u2013',  # ط
    '\u0628': CTRL + '\u201A',  # ب
    '\u0646': CTRL + '\u00A0',  # ن
    '\u0645': CTRL + '\u0178',  # م
    '\u060C': CTRL + '\u00ED',  # ،
    '\u06D4': CTRL + '\u00F3',  # ۔
    '\u064B': CTRL + '\u00C7',  # ً  Fathatan
    '\u064A': CTRL + '\u00B8',  # ي  Arabic Ye
    '\u0610': CTRL + '\u00F8',  # ؐ
    '\u0626': CTRL + '\u00A3',  # ئ
    '\u064C': CTRL + '\u00B5',  # ٌ  Dammatan
    '\u0651': CTRL + '\u00AD',  # ّ  Shadda
    '\u0652': CTRL + '\u00B1',  # ْ  Sukun
    '\uFDFA': CTRL + '\u00F6',  # ﷺ
    '\u0611': CTRL + '\u00AE',  # ؑ
    '\u0691': CTRL + '\u008F',  # ڑ
    '\u0679': CTRL + '\u2026',  # ٹ
    '\u0601': CTRL + '\u00F7',  # ؁
    '\u0657': CTRL + '\u00BE',  # ٗ
    '\u0670': CTRL + '\u00BD',  # ٰ  Superscript Alef
    '\u06C3': CTRL + '\u00B9',  # ۃ
    '\u064F': CTRL + '\u00AC',  # ُ  Damma
    '\u2018': CTRL + '\u00FD',  # '
    '\u2019': CTRL + '\u00FE',  # '
    '\u0614': CTRL + '\u00CF',  # ؔ
    '\u0653': CTRL + '\u00B3',  # ٓ  Maddah
    '\u0635': CTRL + '\u201D',  # ص
    '\u0688': CTRL + '\u0152',  # ڈ
    '\u0656': CTRL + '\u00B0',  # ٖ
    '\u063A': CTRL + '\u2122',  # غ
    '\u062D': CTRL + '\u2030',  # ح
    '\u0636': CTRL + '\u2022',  # ض
    '\u062E': CTRL + '\u0160',  # خ
    '\u0612': CTRL + '\u00E7',  # ؒ
    '\u0630': CTRL + '\u008D',  # ذ
    '\u0698': CTRL + '\u2018',  # ژ
    '\u062B': CTRL + '\u2020',  # ث
    '\u0638': CTRL + '\u2014',  # ظ
    '\u0613': CTRL + '\u00E6',  # ؓ
    '\u0650': CTRL + '\u00AA',  # ِ  Kasra
    '\u064E': CTRL + '\u00AB',  # َ  Fatha
    '\u061F': CTRL + '\u00EE',  # ؟
    '[':      CTRL + '\u00FA',
    ']':      CTRL + '\u00FB',
    '.':      CTRL + '\u00FC',
    '!':      CTRL + '\u00DA',
    ',':      CTRL + '\u00F9',
    '/':      CTRL + '\u00F1',
    ')':      CTRL + '\u00E1',
    '(':      CTRL + '\u00E2',
    ':':      CTRL + '\u00E9',
    ' ':      CTRL + ' ',

    # ══ Arabic-specific letters ══
    '\u0647': CTRL + '\u00A6',  # ه  Arabic He   -> same as ہ in Nastaleeq
    '\u0643': CTRL + '\u0153',  # ك  Arabic Kaf  -> same as ک in Nastaleeq
    '\u0629': CTRL + '\u00B4',  # ة  Teh Marbuta -> unique glyph
    '\u0649': CTRL + '\u00A4',  # ى  Alef Maqsura-> same as ی in Nastaleeq
    '\u0625': CTRL + '\u0081',  # إ Alef hamza below -> same as ا in Nastaleeq
}

_RECOGNIZED_CHARS = set(UNICODE_TO_INPAGE.keys()) | _PRE_REPLACEMENT_CHARS
_URDU_RE   = re.compile(r'[\u0600-\u06FF\u0750-\u077F\uFB50-\uFDFF\uFE70-\uFEFF]')
_LATIN_RE  = re.compile(r'[A-Za-z]')

# ══════════════════════════════════════════════════════════════
#  STEP 1: Symbol Cleanup & Typo Fixes
# ══════════════════════════════════════════════════════════════

def apply_typo_fixes(text: str) -> str:
    for wrong, right in TYPO_FIXES.items():
        text = text.replace(wrong, right)
    return text

def apply_symbol_replacements(text: str) -> tuple[str, Counter]:
    replaced_counts = Counter()
    for sym, replacement in SYMBOL_REPLACEMENTS.items():
        if sym in text:
            count = text.count(sym)
            replaced_counts[sym] += count
            text = text.replace(sym, replacement)
    return text, replaced_counts

# ══════════════════════════════════════════════════════════════
#  STEP 2: English RTL Fix
# ══════════════════════════════════════════════════════════════

def fix_english_for_rtl(line: str) -> str:
    if not _LATIN_RE.search(line):
        return line

    segments = []
    current = ''
    current_type = None

    for ch in line:
        if _URDU_RE.match(ch):
            ch_type = 'urdu'
        elif _LATIN_RE.match(ch) or ch.isdigit():
            ch_type = 'english'
        else:
            ch_type = current_type or 'other'

        if ch_type != current_type and current:
            segments.append((current_type, current))
            current = ch
            current_type = ch_type
        else:
            current += ch
            current_type = ch_type

    if current:
        segments.append((current_type, current))

    result = []
    for seg_type, seg_text in segments:
        if seg_type == 'english' and _LATIN_RE.search(seg_text):
            stripped   = seg_text.strip()
            lead_sp    = seg_text[:len(seg_text) - len(seg_text.lstrip())]
            trail_sp   = seg_text[len(seg_text.rstrip()):]

            lead_p = ''
            trail_p = ''
            tmp = stripped
            while tmp and not tmp[0].isalnum() and tmp[0] != ' ':
                lead_p += tmp[0]; tmp = tmp[1:]
            while tmp and not tmp[-1].isalnum() and tmp[-1] != ' ':
                trail_p = tmp[-1] + trail_p; tmp = tmp[:-1]

            words = tmp.split(' ')
            words.reverse()
            reversed_text = trail_p + ' '.join(words) + lead_p
            result.append(lead_sp + reversed_text + trail_sp)
        else:
            result.append(seg_text)

    return ''.join(result)

# ══════════════════════════════════════════════════════════════
#  STEP 3: Core Unicode -> InPage Conversion
# ══════════════════════════════════════════════════════════════

def convert_line(line: str) -> tuple[str, set]:
    unmapped = set()

    for ch in line:
        if ch in ('\n', '\t', '\r', CTRL):
            continue
        if ch.strip() and ch not in _RECOGNIZED_CHARS:
            unmapped.add(ch)

    for old, new in PRE_REPLACEMENTS:
        line = line.replace(old, new)

    result = []
    for ch in line:
        if ch in UNICODE_TO_INPAGE:
            result.append(UNICODE_TO_INPAGE[ch])
        elif ch in ('\n', '\r', '\t'):
            result.append(ch)
        elif ch == CTRL:
            result.append(ch)
        else:
            result.append(ch)

    return ''.join(result), unmapped

def convert_text(text: str) -> tuple[str, set, int, int]:
    # ── V4.2 CORE FIX: AUTO-SPACE INJECTION ──
    text = re.sub(r'([\u06D2\u06BA])(?=[^\s\x00-\x7F\u060C\u061B\u061F\u06D4])', r'\1 ', text)

    lines = text.split('\n')
    converted_lines = []
    all_unmapped = set()
    total_chars = 0
    unconverted_chars = 0

    for line in lines:
        line = fix_english_for_rtl(line)
        converted, unmapped = convert_line(line)
        converted_lines.append(converted)
        all_unmapped |= unmapped

        for ch in line:
            if ch not in ('\n', '\r', '\t') and ch.strip():
                total_chars += 1
                if ch in unmapped:
                    unconverted_chars += 1

    return '\n'.join(converted_lines), all_unmapped, total_chars, unconverted_chars

# ══════════════════════════════════════════════════════════════
#  FILE READERS
# ══════════════════════════════════════════════════════════════

def read_txt(path: str) -> str:
    for enc in ['utf-8-sig', 'utf-8', 'utf-16', 'utf-16-le', 'utf-16-be', 'cp1256', 'iso-8859-6']:
        try:
            with open(path, 'r', encoding=enc) as f:
                content = f.read()
            print(f"    ✓ Encoding: {enc}")
            return content
        except (UnicodeDecodeError, LookupError):
            continue
    raise ValueError(f"File '{path}' kisi encoding se nahi parhi ja saki.")

def read_docx(path: str) -> str:
    try:
        from docx import Document
    except ImportError:
        print("\n  ✗ python-docx install karo:  pip install python-docx\n")
        sys.exit(1)
    doc = Document(path)
    lines = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            cells = ' | '.join(c.text.strip() for c in row.cells if c.text.strip())
            if cells:
                lines.append(cells)
    return '\n'.join(lines)

# ══════════════════════════════════════════════════════════════
#  REPORT PRINTER
# ══════════════════════════════════════════════════════════════

def print_report(unmapped: set, replaced: Counter, total: int, unconverted: int):
    converted = total - unconverted
    pct = round(converted / max(total, 1) * 100, 2)

    print()
    print("┌─────────────────────────────────────────────┐")
    print("│              CONVERSION REPORT              │")
    print("├─────────────────────────────────────────────┤")
    print(f"│  Total characters    : {total:>8,}             │")
    print(f"│  Successfully mapped : {converted:>8,}  ({pct}%)  │")
    print(f"│  Unmapped (as-is)    : {unconverted:>8,}             │")
    print("├─────────────────────────────────────────────┤")

    if replaced:
        print("│  Symbols replaced:                          │")
        for sym, cnt in replaced.most_common(10):
            name = f"U+{ord(sym):04X}"
            print(f"│    '{sym}' ({name}) x{cnt:<6}                 │")

    if unmapped:
        print("├─────────────────────────────────────────────┤")
        print("│  Still unmapped (check InPage output):      │")
        for ch in sorted(unmapped)[:15]:
            print(f"│    '{ch}'  U+{ord(ch):04X}                        │")
        if len(unmapped) > 15:
            print(f"│    ... aur {len(unmapped)-15} aur characters             │")
    else:
        print("│  ✓ Koi unmapped character nahi!             │")

    print("└─────────────────────────────────────────────┘")

# ══════════════════════════════════════════════════════════════
#  MAIN
# ══════════════════════════════════════════════════════════════

def main():
    print()
    print("╔══════════════════════════════════════════════════════╗")
    print("║    Urdu Novel -> InPage Converter  |  v4.3           ║")
    print("║    v4.3 FIX: Typo Auto-Fixer added for OCR errors    ║")
    print("╚══════════════════════════════════════════════════════╝")
    print(f"  Input   : {INPUT_PATH}")
    print(f"  Output  : {OUTPUT_PATH}")
    print(f"  Symbols : {SYMBOL_ACTION}")
    print("─" * 56)

    if not os.path.exists(INPUT_PATH):
        print(f"\n  ✗ File nahi mili: {INPUT_PATH}")
        sys.exit(1)

    ext = os.path.splitext(INPUT_PATH)[1].lower()
    if ext not in ('.txt', '.docx'):
        print("\n  ✗ Sirf .txt aur .docx support hain.")
        sys.exit(1)

    print("\n  [1/4] File parh raha hoon...")
    text = read_txt(INPUT_PATH) if ext == '.txt' else read_docx(INPUT_PATH)
    
    # Apply Typo Fixes First
    if TYPO_FIXES:
        text = apply_typo_fixes(text)

    char_count = len(text)
    word_count = len(text.split())
    line_count = text.count('\n') + 1
    print(f"    {line_count:,} lines | {word_count:,} words | {char_count:,} chars")

    print("\n  [2/4] Symbols clean kar raha hoon...")
    if SYMBOL_ACTION == 'replace':
        text, replaced_counts = apply_symbol_replacements(text)
        print(f"    {sum(replaced_counts.values()):,} symbols replace hue ({len(replaced_counts)} unique types)")
    elif SYMBOL_ACTION == 'remove':
        replaced_counts = Counter()
        for sym in SYMBOL_REPLACEMENTS:
            if sym in text:
                replaced_counts[sym] += text.count(sym)
                text = text.replace(sym, '')
        print(f"    {sum(replaced_counts.values()):,} symbols remove hue")
    else:
        replaced_counts = Counter()
        print("    Symbols as-is rakhe gaye")

    print("\n  [3/4] Convert + English RTL fix ho raha hai...")
    converted, unmapped, total_chars, unconverted_chars = convert_text(text)
    pct = round((total_chars - unconverted_chars) / max(total_chars, 1) * 100, 1)
    print(f"    {pct}% characters successfully convert hue")

    print("\n  [4/4] File save ho rahi hai...")
    os.makedirs(os.path.dirname(OUTPUT_PATH) or '.', exist_ok=True)
    with open(OUTPUT_PATH, 'w', encoding='utf-8-sig', newline='') as f:
        f.write(converted)
    size_kb = os.path.getsize(OUTPUT_PATH) / 1024
    print(f"    ✓ '{os.path.basename(OUTPUT_PATH)}'  ({size_kb:.1f} KB)")

    print_report(unmapped, replaced_counts, total_chars, unconverted_chars)

    print()
    print("╔══════════════════════════════════════════════════════╗")
    print("║              INPAGE MEIN USE KARNE KA TARIKA         ║")
    print("╠══════════════════════════════════════════════════════╣")
    print(f"║  1. '{os.path.basename(OUTPUT_PATH)}'")
    print("║     Notepad mein kholo                               ║")
    print("║  2. Ctrl+A -> Ctrl+C  (sab copy karo)                ║")
    print("║  3. InPage mein Text Box banao                       ║")
    print("║  4. Ctrl+V paste karo                                ║")
    print("║  5. Font: Noori Nastaleeq ya Jameel Noori            ║")
    print("║  6. Direction: RTL confirm karo                      ║")
    print("║  ─────────────────────────────────────────────────   ║")
    print("║  TIP: Pehle sirf ek chapter paste karo aur check     ║")
    print("║       karo. Theek hai to puri novel karo.            ║")
    print("╚══════════════════════════════════════════════════════╝")
    print()

if __name__ == '__main__':
    main()